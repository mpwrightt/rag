"""
FastAPI endpoints for the agentic RAG system.
"""

import os
import io
import asyncio
import json
import logging
import tempfile
import shutil
from pathlib import Path
import re
import glob
from contextlib import asynccontextmanager
from typing import Dict, Any, List, Optional, Tuple
from datetime import datetime
import uuid

from fastapi import FastAPI, HTTPException, Request, Depends, UploadFile, File, Form
from fastapi.responses import StreamingResponse, Response, JSONResponse, HTMLResponse
from fastapi.middleware.cors import CORSMiddleware
from fastapi.middleware.gzip import GZipMiddleware
import uvicorn
from dotenv import load_dotenv
import csv

from .agent import rag_agent, AgentDependencies
from .enhanced_retrieval import EnhancedRetriever

from .context import get_current_search_results, clear_search_results, register_retrieval_listener, unregister_retrieval_listener, emit_retrieval_event
from .db_utils import (
    initialize_database,
    close_database,
    get_session,
    add_message,
    get_session_messages,
    test_connection,
    delete_document,
    list_documents,
    get_document,
    get_graph_statistics,
    search_facts,
    search_facts_websearch,
    # Incremental update helpers
    update_document_metadata_only,
    update_chunk_metadata_batch,
    add_document_tags,
    # Collections
    list_collections_db,
    create_collection_db,
    update_collection_db,
    delete_collection_db,
    add_documents_to_collection_db,
    remove_document_from_collection_db,
    list_collection_documents_db,
    update_document_collections_db,
    # Proposals
    create_proposal_db,
    get_proposal_db,
    list_proposals_db,
    create_proposal_version_db,
    get_latest_proposal_version_db,
    get_proposal_version_db,
    list_proposal_versions_db,
    update_proposal_db,
    list_proposal_documents_db,
    add_section_feedback_db,
    get_top_rated_section_examples_db,
    # Summary jobs
    create_summary_job,
    update_summary_job_status,
    set_summary_job_result,
    get_summary_job,
    cancel_summary_job,
    is_summary_job_cancelled,
)
from .graph_utils import (
    initialize_graph,
    close_graph,
    test_graph_connection,
    get_entity_relationships as kg_get_entity_relationships,
    search_knowledge_graph as kg_search_knowledge_graph,
)
from .query_processor import QueryProcessor
from .models import (
    ChatRequest,
    ChatResponse,
    SearchRequest,
    SearchResponse,
    StreamDelta,
    ErrorResponse,
    HealthStatus,
    ToolCall,
    SourceResult,
    ChunkResult,
    GraphSearchResult,
    RealTimeMetrics,
    ChatMetrics,
    DocumentUsageStats,
    UserEngagementMetrics,
    AnalyticsDashboardResponse,
    DocumentMetadataUpdateRequest,
    ChunkMetadataBatchUpdateRequest,
    AddTagsRequest,
    UpdateClassificationRequest,
    UpdateCollectionsRequest,
)
from .tools import (
    vector_search_tool,
    graph_search_tool,
    hybrid_search_tool,
    list_documents_tool,
    get_document_tool,
    VectorSearchInput,
    GraphSearchInput,
    HybridSearchInput,
    DocumentListInput,
    DocumentInput
)

# Analytics tracker
from .analytics import analytics_tracker
from .question_generator import QuestionGenerator
from .proposal_analyzer import analyze_example_text

# Load environment variables
load_dotenv()

logger = logging.getLogger(__name__)

# Import ingestion pipeline components
try:
    import sys
    import os
    # Add the parent directory to sys.path to allow imports
    sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
    from ingestion.ingest import DocumentIngestionPipeline
    from ingestion.chunker import ChunkingConfig, create_chunker
    from ingestion.embedder import create_embedder
    from ingestion.graph_builder import create_graph_builder
    from .models import IngestionConfig
    from ingestion.converters import convert_to_markdown
    INGESTION_AVAILABLE = True
except ImportError as e:
    logger.warning(f"Ingestion pipeline not available: {e}")
    INGESTION_AVAILABLE = False

# Import 3rd party integrations
try:
    from integrations import (
        GoogleDriveIntegration, 
        DropboxIntegration, 
        OneDriveIntegration,
        IntegrationConfig,
        create_google_drive_integration,
        create_dropbox_integration, 
        create_onedrive_integration
    )
    INTEGRATIONS_AVAILABLE = True
except ImportError as e:
    logger.warning(f"3rd party integrations not available: {e}")
    INTEGRATIONS_AVAILABLE = False
    
    # Create dummy classes to prevent unbound variable errors
    class GoogleDriveIntegration:
        def __init__(self, *args, **kwargs): pass
    class DropboxIntegration:
        def __init__(self, *args, **kwargs): pass
    class OneDriveIntegration:
        def __init__(self, *args, **kwargs): pass
    class IntegrationConfig:
        def __init__(self, *args, **kwargs): pass
    
    def create_google_drive_integration(*args, **kwargs): return None
    def create_dropbox_integration(*args, **kwargs): return None
    def create_onedrive_integration(*args, **kwargs): return None

# Application configuration
APP_ENV = os.getenv("APP_ENV", "development")
APP_HOST = os.getenv("APP_HOST", "0.0.0.0")
APP_PORT = int(os.getenv("APP_PORT", 8000))
LOG_LEVEL = os.getenv("LOG_LEVEL", "INFO")

# Configure logging
logging.basicConfig(
    level=getattr(logging, LOG_LEVEL.upper()),
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s"
)

# Set debug level for our module during development
if APP_ENV == "development":
    logger.setLevel(logging.DEBUG)


@asynccontextmanager
async def lifespan(app: FastAPI):
    """Lifespan context manager for FastAPI app."""
    # Startup
    logger.info("Starting up agentic RAG API...")
    
    try:
        # Initialize database connections (tolerate failure and continue in degraded mode)
        try:
            await initialize_database()
            logger.info("Database initialized")
        except Exception as e:
            logger.exception("Database initialization failed: %s", e)

        # Initialize graph database (tolerate failure)
        try:
            await initialize_graph()
            logger.info("Graph database initialized")
        except Exception as e:
            logger.exception("Graph initialization failed: %s", e)

        # Test connections (report status but don't block startup)
        try:
            db_ok = await test_connection()
        except Exception as e:
            logger.exception("Database test failed during startup: %s", e)
            db_ok = False
        try:
            graph_ok = await test_graph_connection()
        except Exception as e:
            logger.exception("Graph test failed during startup: %s", e)
            graph_ok = False

        if not db_ok:
            logger.error("Database connection failed")
        if not graph_ok:
            logger.error("Graph database connection failed")

        # Initialize question generator here (lifespan startup)
        global question_generator
        # Accept multiple env var names for compatibility with different configs
        gemini_api_key = (
            os.getenv("GEMINI_API_KEY")
            or os.getenv("GOOGLE_API_KEY")
            or os.getenv("LLM_API_KEY")
        )
        logger.info(
            "Looking for API keys - GEMINI_API_KEY: %s, GOOGLE_API_KEY: %s, LLM_API_KEY: %s",
            "found" if os.getenv("GEMINI_API_KEY") else "not found",
            "found" if os.getenv("GOOGLE_API_KEY") else "not found",
            "found" if os.getenv("LLM_API_KEY") else "not found",
        )
        if gemini_api_key:
            try:
                question_generator = QuestionGenerator(gemini_api_key)
                logger.info("Question generator initialized successfully with Gemini")
            except Exception as e:
                logger.error(f"Failed to initialize question generator: {e}")
                question_generator = None
        else:
            logger.warning("GEMINI_API_KEY / GOOGLE_API_KEY / LLM_API_KEY not found, question generation will be disabled")
        
        logger.info("Agentic RAG API startup complete")
        
    except Exception as e:
        # Never fail the app startup; continue in degraded mode so /health can report status
        logger.exception("Startup encountered unexpected error; continuing in degraded mode: %s", e)
    
    yield
    
    # Shutdown
    logger.info("Shutting down agentic RAG API...")
    
    try:
        await close_database()
        await close_graph()
        logger.info("Connections closed")
    except Exception as e:
        logger.error(f"Shutdown error: {e}")


# Create FastAPI app
app = FastAPI(
    title="Agentic RAG with Knowledge Graph",
    description="AI agent combining vector search and knowledge graph for tech company analysis",
    version="0.1.0",
    lifespan=lifespan
)

# Add middleware with flexible CORS
# Configure CORS with environment-driven origins. Using wildcard with credentials can be problematic on some hosts.
allowed_origins_env = os.getenv("ALLOWED_ORIGINS") or os.getenv("CORS_ALLOW_ORIGINS") or ""
allowed_origins = [o.strip() for o in allowed_origins_env.split(",") if o.strip()]

# In production, default to allowing datadiver.app domains via regex if no explicit origins are set
default_origin_regex = None
try:
    if os.getenv("APP_ENV", "development").lower() != "development" and not allowed_origins:
        default_origin_regex = r"https://(.+\.)?datadiver\.app$"
except Exception:
    default_origin_regex = None

app.add_middleware(
    CORSMiddleware,
    allow_origins=allowed_origins or ([] if default_origin_regex else ["*"]),
    allow_origin_regex=os.getenv("CORS_ALLOW_ORIGIN_REGEX") or default_origin_regex,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"]
)

app.add_middleware(GZipMiddleware, minimum_size=1000)

# Initialize question generator
question_generator = None

# In-memory store for multipart upload sessions (best-effort, non-persistent)
UPLOAD_SESSIONS: Dict[str, Dict[str, Any]] = {}

# In-memory ingest job tracker (ephemeral)
INGEST_JOBS: Dict[str, Dict[str, Any]] = {}

# In-memory proposal helpers (ephemeral, process-lifetime)
# - PROPOSAL_STYLE_HINTS: stores example analysis such as style_prompt and inferred sections
# - PROPOSAL_DRAFT_TEXTS: stores raw draft text uploaded by user
PROPOSAL_STYLE_HINTS: Dict[str, Dict[str, Any]] = {}
PROPOSAL_DRAFT_TEXTS: Dict[str, str] = {}


# Helper functions for agent execution
async def get_or_create_session(request: ChatRequest) -> str:
    """Get existing session or create new one."""
    # If client supplies a session_id, try to use/validate it but fall back gracefully
    if request.session_id:
        try:
            session = await get_session(request.session_id)
            if session:
                return request.session_id
        except Exception:
            pass
    # Create a new session (lazy import to avoid startup ImportError)
    try:
        from .db_utils import create_session as _create_session
        new_session_id = await _create_session(user_id=None, metadata={"source": "api"})
        return new_session_id
    except Exception as e:
        logger.warning("create_session unavailable, generating ephemeral session id: %s", e)
        return str(uuid.uuid4())


async def get_conversation_context(
    session_id: str,
    max_messages: int = 10
) -> List[Dict[str, str]]:
    """
    Get recent conversation context.
    
    Args:
        session_id: Session ID
        max_messages: Maximum number of messages to retrieve
    
    Returns:
        List of messages
    """
    messages = await get_session_messages(session_id, limit=max_messages)
    
    return [
        {
            "role": msg["role"],
            "content": msg["content"]
        }
        for msg in messages
    ]


def extract_tool_calls(result) -> List[ToolCall]:
    """
    Extract tool calls from Pydantic AI result.
    
    Args:
        result: Pydantic AI result object
    
    Returns:
        List of ToolCall objects
    """
    tools_used = []
    
    try:
        # Get all messages from the result
        messages = result.all_messages()
        
        for message in messages:
            if hasattr(message, 'parts'):
                for part in message.parts:
                    # Check if this is a tool call part
                    if part.__class__.__name__ == 'ToolCallPart':
                        try:
                            # Debug logging to understand structure
                            logger.debug(f"ToolCallPart attributes: {dir(part)}")
                            logger.debug(f"ToolCallPart content: tool_name={getattr(part, 'tool_name', None)}")
                            
                            # Extract tool information safely
                            tool_name = str(part.tool_name) if hasattr(part, 'tool_name') else 'unknown'
                            
                            # Get args - the args field is a JSON string in Pydantic AI
                            tool_args = {}
                            if hasattr(part, 'args') and part.args is not None:
                                if isinstance(part.args, str):
                                    # Args is a JSON string, parse it
                                    try:
                                        import json
                                        tool_args = json.loads(part.args)
                                        logger.debug(f"Parsed args from JSON string: {tool_args}")
                                    except json.JSONDecodeError as e:
                                        logger.debug(f"Failed to parse args JSON: {e}")
                                        tool_args = {}
                                elif isinstance(part.args, dict):
                                    tool_args = part.args
                                    logger.debug(f"Args already a dict: {tool_args}")
                            
                            # Alternative: use args_as_dict method if available
                            if hasattr(part, 'args_as_dict'):
                                try:
                                    tool_args = part.args_as_dict()
                                    logger.debug(f"Got args from args_as_dict(): {tool_args}")
                                except:
                                    pass
                            
                            # Get tool call ID
                            tool_call_id = None
                            if hasattr(part, 'tool_call_id'):
                                tool_call_id = str(part.tool_call_id) if part.tool_call_id else None
                            
                            # Create ToolCall with explicit field mapping
                            tool_call_data = {
                                "tool_name": tool_name,
                                "args": tool_args,
                                "tool_call_id": tool_call_id
                            }
                            logger.debug(f"Creating ToolCall with data: {tool_call_data}")
                            tools_used.append(ToolCall(**tool_call_data))
                        except Exception as e:
                            logger.debug(f"Failed to parse tool call part: {e}")
                            continue
    except Exception as e:
        logger.warning(f"Failed to extract tool calls: {e}")
    
    return tools_used


async def save_conversation_turn(
    session_id: str,
    user_message: str,
    assistant_message: str,
    metadata: Optional[Dict[str, Any]] = None
):
    """
    Save a conversation turn to the database.
    
    Args:
        session_id: Session ID
        user_message: User's message
        assistant_message: Assistant's response
        metadata: Optional metadata
    """
    # Save user message
    await add_message(
        session_id=session_id,
        role="user",
        content=user_message,
        metadata=metadata or {}
    )
    
    # Save assistant message
    await add_message(
        session_id=session_id,
        role="assistant",
        content=assistant_message,
        metadata=metadata or {}
    )


def convert_chunks_to_sources(chunks: List[ChunkResult]) -> List[SourceResult]:
    """Convert ChunkResult objects to SourceResult objects for frontend."""
    sources = []
    seen_combinations = set()
    
    for chunk in chunks:
        # Create unique identifier to avoid duplicates
        combo_id = f"{chunk.document_source}:{chunk.chunk_id}"
        if combo_id not in seen_combinations:
            seen_combinations.add(combo_id)
            sources.append(SourceResult(
                filename=chunk.document_source,
                chunk_id=chunk.chunk_id,
                relevance_score=chunk.score,
                document_title=chunk.document_title
            ))
    
    return sources


async def build_live_graph_payload(
    message: str,
    tools_used: List[ToolCall],
    max_expand_entities: int = 3,
    default_depth: int = 2,
) -> Optional[Dict[str, Any]]:
    """
    Build a live knowledge graph payload using the PostgreSQL graph layer.
    Prefers explicit get_entity_relationships tool args; falls back to graph_search.
    """
    try:
        # Prefer explicit entity from get_entity_relationships tool usage
        entity_name: Optional[str] = None
        entity_depth: int = default_depth
        graph_query: Optional[str] = None

        for t in tools_used or []:
            tool_name = (getattr(t, "tool_name", "") or "").lower()
            args = getattr(t, "args", {}) or {}
            if "get_entity_relationships" in tool_name:
                entity_name = args.get("entity_name") or args.get("entity")
                try:
                    if isinstance(args.get("depth"), int):
                        entity_depth = max(1, min(5, int(args["depth"])))
                except Exception:
                    pass
            elif "graph_search" in tool_name and not graph_query:
                graph_query = args.get("query")

        # If we have an explicit entity, build graph from relationships
        if entity_name:
            rel_result = await kg_get_entity_relationships(entity_name, depth=entity_depth)
            relationships = rel_result.get("relationships", []) if isinstance(rel_result, dict) else []

            nodes_map: Dict[str, Dict[str, Any]] = {}
            edges: List[Dict[str, Any]] = []
            for r in relationships:
                s_name = r.get("source_name")
                s_type = r.get("source_type")
                t_name = r.get("target_name")
                t_type = r.get("target_type")
                rel = r.get("relationship_type") or "related_to"
                if not (s_name and s_type and t_name and t_type):
                    continue
                s_id = f"{s_name}|{s_type}"
                t_id = f"{t_name}|{t_type}"
                if s_id not in nodes_map:
                    nodes_map[s_id] = {"id": s_id, "label": s_name, "type": s_type.capitalize() if isinstance(s_type, str) else s_type}
                if t_id not in nodes_map:
                    nodes_map[t_id] = {"id": t_id, "label": t_name, "type": t_type.capitalize() if isinstance(t_type, str) else t_type}
                edges.append({"source": s_id, "target": t_id, "relationship": rel})

            # Ensure the central node exists
            if entity_name and not any(entity_name == n.get("label") for n in nodes_map.values()):
                center_id = f"{entity_name}|Entity"
                nodes_map[center_id] = {"id": center_id, "label": entity_name, "type": "Entity"}

            return {"nodes": list(nodes_map.values()), "edges": edges}

        # Fallback: use graph search on the query or message to seed entities, then expand shallow relationships
        search_q = graph_query or (message or "").strip()
        if not search_q:
            return None

        facts = await kg_search_knowledge_graph(search_q)
        # Extract top unique node names/types from facts
        seen = set()
        seed_entities: List[Tuple[str, str]] = []
        for f in facts or []:
            name = f.get("node_name")
            ntype = f.get("node_type")
            key = (name, ntype)
            if name and ntype and key not in seen:
                seen.add(key)
                seed_entities.append(key)
            if len(seed_entities) >= max_expand_entities:
                break

        nodes_map: Dict[str, Dict[str, Any]] = {}
        edge_set = set()

        # Add seed nodes
        for name, ntype in seed_entities:
            nid = f"{name}|{ntype}"
            nodes_map[nid] = {"id": nid, "label": name, "type": ntype.capitalize() if isinstance(ntype, str) else ntype}

        # Expand relationships shallowly
        for name, _ntype in seed_entities:
            try:
                rels = await kg_get_entity_relationships(name, depth=1)
            except Exception:
                continue
            for r in (rels.get("relationships", []) if isinstance(rels, dict) else []):
                s_name = r.get("source_name")
                s_type = r.get("source_type")
                t_name = r.get("target_name")
                t_type = r.get("target_type")
                rel = r.get("relationship_type") or "related_to"
                if not (s_name and s_type and t_name and t_type):
                    continue
                s_id = f"{s_name}|{s_type}"
                t_id = f"{t_name}|{t_type}"
                if s_id not in nodes_map:
                    nodes_map[s_id] = {"id": s_id, "label": s_name, "type": s_type.capitalize() if isinstance(s_type, str) else s_type}
                if t_id not in nodes_map:
                    nodes_map[t_id] = {"id": t_id, "label": t_name, "type": t_type.capitalize() if isinstance(t_type, str) else t_type}
                edge_key = (s_id, rel, t_id)
                if edge_key not in edge_set:
                    edge_set.add(edge_key)

        edges = [
            {"source": s, "relationship": rel, "target": t}
            for (s, rel, t) in edge_set
        ]

        if nodes_map or edges:
            return {"nodes": list(nodes_map.values()), "edges": edges}

        return None
    except Exception as e:
        logger.debug(f"Failed to build live graph payload: {e}")
        return None


async def execute_agent(
    message: str,
    session_id: str,
    user_id: Optional[str] = None,
    save_conversation: bool = True,
    search_preferences: Optional[Dict[str, Any]] = None,
) -> tuple[str, List[ToolCall], List[SourceResult]]:
    """
    Execute the agent with a message.
    
    Args:
        message: User message
        session_id: Session ID
        user_id: Optional user ID
        save_conversation: Whether to save the conversation
    
    Returns:
        Tuple of (agent response, tools used, sources)
    """
    try:
        # Clear previous search results
        clear_search_results()
        
        # Create dependencies (include search scoping preferences if provided)
        deps = AgentDependencies(
            session_id=session_id,
            user_id=user_id,
            search_preferences=search_preferences or {},
        )
        
        # Get conversation context
        context = await get_conversation_context(session_id)
        
        # Build prompt with context
        full_prompt = message
        if context:
            context_str = "\n".join([
                f"{msg['role']}: {msg['content']}"
                for msg in context[-6:]  # Last 3 turns
            ])
            full_prompt = f"Previous conversation:\n{context_str}\n\nCurrent question: {message}"
        
        # Run the agent
        result = await rag_agent.run(full_prompt, deps=deps)
        
        response = result.data
        tools_used = extract_tool_calls(result)
        
        # Convert captured search results to sources
        sources = convert_chunks_to_sources(get_current_search_results())
        
        # Save conversation if requested
        if save_conversation:
            await save_conversation_turn(
                session_id=session_id,
                user_message=message,
                assistant_message=response,
                metadata={
                    "user_id": user_id,
                    "tool_calls": len(tools_used),
                    "sources_found": len(sources)
                }
            )
        
        return response, tools_used, sources
        
    except Exception as e:
        logger.error(f"Agent execution failed: {e}")
        error_response = f"I encountered an error while processing your request: {str(e)}"
        
        if save_conversation:
            await save_conversation_turn(
                session_id=session_id,
                user_message=message,
                assistant_message=error_response,
                metadata={"error": str(e)}
            )
        
        return error_response, [], []


# API Endpoints

@app.get("/collections")
async def api_list_collections(
    page: int = 1,
    per_page: int = 20,
    search: Optional[str] = None,
    created_by: Optional[str] = None,
    workspace_id: Optional[str] = None,
    is_shared: Optional[bool] = None,
):
    try:
        offset = (page - 1) * per_page
        collections, total = await list_collections_db(
            limit=per_page,
            offset=offset,
            search=search,
            created_by=created_by,
            workspace_id=workspace_id,
            is_shared=is_shared,
        )
        return {"collections": collections, "total": total, "page": page, "per_page": per_page}
    except Exception as e:
        logger.exception("Failed to list collections: %s", e)
        raise HTTPException(status_code=500, detail="Failed to list collections")


@app.post("/collections")
async def api_create_collection(request: Request):
    try:
        body = await request.json()
        created = await create_collection_db(
            name=body.get("name"),
            description=body.get("description"),
            color=body.get("color") or "#6366f1",
            icon=body.get("icon") or "folder",
            is_shared=bool(body.get("is_shared", False)),
            created_by=request.headers.get("x-user-id"),
            workspace_id=request.headers.get("x-workspace-id"),
            metadata=body.get("metadata") or {},
        )
        return created
    except Exception as e:
        logger.exception("Failed to create collection: %s", e)
        raise HTTPException(status_code=500, detail="Failed to create collection")


@app.patch("/collections/{collection_id}")
async def api_update_collection(collection_id: str, request: Request):
    try:
        body = await request.json()
        updated = await update_collection_db(
            collection_id,
            name=body.get("name"),
            description=body.get("description"),
            color=body.get("color"),
            icon=body.get("icon"),
            is_shared=body.get("is_shared"),
            metadata=body.get("metadata"),
        )
        if not updated:
            raise HTTPException(status_code=404, detail="Collection not found")
        return updated
    except HTTPException:
        raise
    except Exception as e:
        logger.exception("Failed to update collection: %s", e)
        raise HTTPException(status_code=500, detail="Failed to update collection")


@app.delete("/collections/{collection_id}")
async def api_delete_collection(collection_id: str):
    try:
        ok = await delete_collection_db(collection_id)
        if not ok:
            raise HTTPException(status_code=404, detail="Collection not found")
        return {"success": True}
    except HTTPException:
        raise
    except Exception as e:
        logger.exception("Failed to delete collection: %s", e)
        raise HTTPException(status_code=500, detail="Failed to delete collection")


@app.get("/collections/{collection_id}/documents")
async def api_list_collection_documents(collection_id: str, page: int = 1, per_page: int = 50):
    try:
        docs, total = await list_collection_documents_db(collection_id, limit=per_page, offset=(page - 1) * per_page)
        return {"documents": docs, "total": total, "page": page, "per_page": per_page}
    except Exception as e:
        logger.exception("Failed to list collection documents: %s", e)
        raise HTTPException(status_code=500, detail="Failed to list collection documents")


@app.post("/collections/{collection_id}/documents")
async def api_add_documents_to_collection(collection_id: str, request: Request):
    try:
        body = await request.json()
        ids = body.get("document_ids") or []
        if not isinstance(ids, list) or not ids:
            raise HTTPException(status_code=400, detail="document_ids array required")
        try:
            added = await add_documents_to_collection_db(collection_id, ids, added_by=request.headers.get("x-user-id"))
        except ValueError as ve:
            # Surface clearer API errors
            if str(ve) == "collection_not_found":
                raise HTTPException(status_code=404, detail="Collection not found")
            raise
        return {"added": added}
    except HTTPException:
        raise
    except Exception as e:
        logger.exception("Failed to add documents to collection: %s", e)
        # Surface underlying error for debugging during integration
        raise HTTPException(status_code=500, detail=str(e))


# -------------------------
# Proposal Management Endpoints
# -------------------------

@app.get("/proposals")
async def api_list_proposals(page: int = 1, per_page: int = 50):
    """List proposals with pagination for dashboard."""
    try:
        offset = (page - 1) * per_page
        items = await list_proposals_db(limit=per_page, offset=offset)
        total = len(items) if isinstance(items, list) else 0
        return {"proposals": items, "total": total, "page": page, "per_page": per_page}
    except Exception as e:
        logger.exception("Failed to list proposals: %s", e)
        raise HTTPException(status_code=500, detail="Failed to list proposals")


@app.post("/proposals")
async def api_create_proposal(request: Request):
    """Create a new proposal and return the created record."""
    try:
        body = await request.json()
        created = await create_proposal_db(
            title=(body.get("title") or "Untitled Proposal").strip(),
            client_fields=body.get("client_fields") or {},
            project_fields=body.get("project_fields") or {},
            status=body.get("status") or "draft",
            metadata=body.get("metadata") or {},
            created_by=request.headers.get("x-user-id"),
        )
        return created
    except Exception as e:
        logger.exception("Failed to create proposal: %s", e)
        raise HTTPException(status_code=500, detail="Failed to create proposal")


@app.get("/proposals/{proposal_id}")
async def api_get_proposal(proposal_id: str):
    """Fetch a single proposal by ID."""
    try:
        prop = await get_proposal_db(proposal_id)
        if not prop:
            raise HTTPException(status_code=404, detail="Proposal not found")
        return prop
    except HTTPException:
        raise
    except Exception as e:
        logger.exception("Failed to get proposal: %s", e)
        raise HTTPException(status_code=500, detail="Failed to get proposal")


@app.get("/proposals/{proposal_id}/regulatory")
async def api_list_proposal_regulatory_documents(proposal_id: str):
    """List regulatory documents attached to a proposal."""
    try:
        docs = await list_proposal_documents_db(proposal_id, source_type="regulatory")
        return {"documents": docs}
    except Exception as e:
        logger.exception("Failed to list proposal regulatory documents: %s", e)
        raise HTTPException(status_code=500, detail="Failed to list regulatory documents")


@app.post("/proposals/{proposal_id}/regulatory/upload")
async def api_upload_proposal_regulatory(proposal_id: str, file: UploadFile = File(...), fast: int = 0):
    """Upload and ingest a regulatory document scoped to a proposal.

    Query param 'fast=1' skips graph building to reduce processing time.
    """
    if not INGESTION_AVAILABLE:
        raise HTTPException(status_code=503, detail="Document ingestion pipeline is not available")
    try:
        # Validate type
        allowed_extensions = {'.txt', '.md', '.pdf', '.doc', '.docx', '.xls', '.xlsx', '.csv', '.tsv'}
        ext = os.path.splitext(file.filename or "")[1].lower()
        if ext not in allowed_extensions:
            raise HTTPException(status_code=400, detail=f"File type {ext} not supported")

        # Stage upload -> convert to markdown
        with tempfile.TemporaryDirectory() as temp_dir:
            raw_path = Path(temp_dir) / (file.filename or "upload.bin")
            size_bytes = 0
            with open(raw_path, 'wb') as f:
                while True:
                    chunk = await file.read(8 * 1024 * 1024)
                    if not chunk:
                        break
                    size_bytes += len(chunk)
                    f.write(chunk)

            try:
                md_text, conv_meta = convert_to_markdown(str(raw_path))
            except Exception as e:
                logger.exception("Regulatory conversion failed: %s", e)
                raise HTTPException(status_code=500, detail="Failed to convert file to markdown")
            if not md_text or not md_text.strip():
                raise HTTPException(status_code=400, detail="No extractable text content found")

            md_path = Path(temp_dir) / f"{Path(file.filename or 'upload').stem}.md"
            with open(md_path, 'w', encoding='utf-8') as f_md:
                f_md.write(md_text)

            # Configure ingestion
            config = IngestionConfig(
                chunk_size=int(os.getenv("CHUNK_SIZE", "800")),
                chunk_overlap=int(os.getenv("CHUNK_OVERLAP", "150")),
                max_chunk_size=int(os.getenv("MAX_CHUNK_SIZE", "1500")),
                use_semantic_splitting=os.getenv("USE_SEMANTIC_SPLITTING", "1") not in ("0", "false", "False"),
                extract_entities=os.getenv("EXTRACT_ENTITIES", "1") not in ("0", "false", "False"),
                skip_graph_building=bool(fast),
            )

            pipeline = DocumentIngestionPipeline(
                config=config,
                documents_folder=str(temp_dir),
                clean_before_ingest=False,
                default_metadata={
                    "proposal_id": proposal_id,
                    "proposal_source_type": "regulatory",
                },
            )

            results = await pipeline.ingest_documents()
            ok = [r for r in results if r.document_id]
            return {
                "success": bool(ok),
                "documents_processed": len(ok),
                "filename": file.filename,
                "size": size_bytes,
            }
    except HTTPException:
        raise
    except Exception as e:
        logger.exception("Failed to upload regulatory document: %s", e)
        raise HTTPException(status_code=500, detail="Failed to upload regulatory document")


@app.delete("/proposals/{proposal_id}/regulatory/{document_id}")
async def api_delete_proposal_regulatory(proposal_id: str, document_id: str):
    """Remove a regulatory document from a proposal by deleting the document."""
    try:
        removed = await delete_document(document_id)
        if not removed:
            raise HTTPException(status_code=404, detail="Document not found")
        return {"removed": True}
    except HTTPException:
        raise
    except Exception as e:
        logger.exception("Failed to delete regulatory document: %s", e)
        raise HTTPException(status_code=500, detail="Failed to delete regulatory document")


@app.post("/proposals/{proposal_id}/example/upload")
async def api_upload_proposal_example(proposal_id: str, file: UploadFile = File(...)):
    """Analyze an example proposal and return structure/style hints."""
    try:
        # Convert to markdown/plaintext
        with tempfile.TemporaryDirectory() as temp_dir:
            raw_path = Path(temp_dir) / (file.filename or "example.bin")
            with open(raw_path, 'wb') as f:
                while True:
                    chunk = await file.read(4 * 1024 * 1024)
                    if not chunk:
                        break
                    f.write(chunk)
            try:
                md_text, _ = convert_to_markdown(str(raw_path))
            except Exception:
                # Fallback: try to read as text
                md_text = raw_path.read_text(encoding='utf-8', errors='ignore')

        analysis = analyze_example_text(md_text or "")
        # Store in ephemeral memory for later generation prompts
        try:
            hints = analysis or {}
            raw_text = (md_text or "").strip()
            # Detect common bullet marker used in the example (prefer non-markdown symbols)
            bullet_candidates = ["", "•", "●", "▪", "–", "-", "*", "·"]
            bullet_counts = {}
            for line in raw_text.splitlines():
                s = line.strip()
                for b in bullet_candidates:
                    if s.startswith(b + " "):
                        bullet_counts[b] = bullet_counts.get(b, 0) + 1
            bullet_marker = None
            if bullet_counts:
                # pick the most frequent bullet
                bullet_marker = sorted(bullet_counts.items(), key=lambda x: x[1], reverse=True)[0][0]
            # Detect Task heading pattern and AOC pattern
            has_task_pattern = bool(re.search(r"(?mi)^Task\s+\d+\s*:\s*", raw_text))
            has_aoc_pattern = bool(re.search(r"(?mi)^AOC[-\s]*\d+", raw_text))
            has_re_line = bool(re.search(r"(?mi)^RE:\s*", raw_text))
            has_salutation = bool(re.search(r"(?mi)^Dear\s+", raw_text))
            hints["formatting"] = {
                "bullet_marker": bullet_marker,
                "has_task_pattern": has_task_pattern,
                "has_aoc_pattern": has_aoc_pattern,
                "has_re_line": has_re_line,
                "has_salutation": has_salutation,
            }
            hints["raw_text"] = raw_text[:25000]  # keep a bounded amount
            PROPOSAL_STYLE_HINTS[proposal_id] = hints
        except Exception:
            pass
        return {"analysis": analysis}
    except Exception as e:
        logger.exception("Failed to analyze example proposal: %s", e)
        raise HTTPException(status_code=500, detail="Failed to analyze example proposal")


@app.post("/proposals/{proposal_id}/draft/upload")
async def api_upload_proposal_draft(proposal_id: str, file: UploadFile = File(...)):
    """Capture draft text for a proposal (for auto-fill later)."""
    try:
        with tempfile.TemporaryDirectory() as temp_dir:
            raw_path = Path(temp_dir) / (file.filename or "draft.bin")
            with open(raw_path, 'wb') as f:
                while True:
                    chunk = await file.read(4 * 1024 * 1024)
                    if not chunk:
                        break
                    f.write(chunk)
            try:
                md_text, _ = convert_to_markdown(str(raw_path))
            except Exception:
                md_text = raw_path.read_text(encoding='utf-8', errors='ignore')
        # Store for later auto-fill/use in prompts (ephemeral)
        try:
            PROPOSAL_DRAFT_TEXTS[proposal_id] = md_text or ""
        except Exception:
            pass
        return {"characters": len(md_text or "")}
    except Exception as e:
        logger.exception("Failed to upload draft: %s", e)
        raise HTTPException(status_code=500, detail="Failed to upload draft")


@app.post("/proposals/{proposal_id}/versions")
async def api_create_proposal_version(proposal_id: str, request: Request):
    """Create a version snapshot for a proposal (sections/citations or HTML)."""
    try:
        body = await request.json()
        version = await create_proposal_version_db(
            proposal_id=proposal_id,
            html=body.get("html"),
            sections=body.get("sections"),
            citations=body.get("citations"),
        )
        return version
    except Exception as e:
        logger.exception("Failed to create proposal version: %s", e)
        raise HTTPException(status_code=500, detail="Failed to create proposal version")


@app.post("/proposals/{proposal_id}/versions/{version_id}/sections/{section_key}/feedback")
async def api_submit_section_feedback(proposal_id: str, version_id: str, section_key: str, request: Request):
    """Record thumbs up/down feedback for a specific section on a proposal version.

    Body JSON:
      - rating: 1 (thumbs up) or -1 (thumbs down)
    """
    try:
        body = await request.json()
        try:
            rating = int(body.get("rating"))
        except Exception:
            raise HTTPException(status_code=400, detail="rating must be 1 or -1")
        if rating not in (1, -1):
            raise HTTPException(status_code=400, detail="rating must be 1 or -1")

        user_id = request.headers.get("x-user-id")
        updated = await add_section_feedback_db(
            proposal_id=proposal_id,
            version_id=version_id,
            section_key=section_key,
            rating=rating,
            user_id=user_id,
        )
        return {"ok": True, "section_ratings": updated}
    except HTTPException:
        raise
    except Exception as e:
        logger.exception("Failed to submit section feedback: %s", e)
        raise HTTPException(status_code=500, detail="Failed to submit section feedback")


@app.post("/proposals/{proposal_id}/validate")
async def api_validate_proposal(proposal_id: str):
    """Validate latest proposal version. For now, returns a simple OK stub."""
    try:
        # Placeholder validation: ensure a version exists
        latest = await get_latest_proposal_version_db(proposal_id)
        if not latest:
            return {"status": "warnings", "warnings": ["No versions saved yet"], "errors": []}
        return {"status": "ok", "warnings": [], "errors": []}
    except Exception as e:
        logger.exception("Failed to validate proposal: %s", e)
        return {"status": "errors", "warnings": [], "errors": [str(e)]}


def _build_proposal_html(sections: Any, title: str = "Proposal") -> str:
    try:
        parts = [
            "<html>",
            "<head>",
            f"<meta charset='utf-8'><title>{title}</title>",
            "<style>body{font-family:system-ui,-apple-system,Segoe UI,Roboto,Helvetica,Arial,sans-serif;margin:24px;} h2{margin-top:24px;border-bottom:1px solid #eee;padding-bottom:4px;} .section{margin-bottom:16px;}</style>",
            "</head>",
            "<body>",
            f"<h1>{title}</h1>",
        ]
        for s in (sections or []):
            stitle = (s.get("title") or "Section") if isinstance(s, dict) else "Section"
            scontent = (s.get("content") or "") if isinstance(s, dict) else ""
            parts.append(f"<div class='section'><h2>{stitle}</h2>\n<div>{scontent}</div></div>")
        parts.append("</body></html>")
        return "\n".join(parts)
    except Exception:
        return "<html><body><h1>Proposal</h1><p>No content.</p></body></html>"


@app.get("/proposals/{proposal_id}/export")
async def api_export_proposal(proposal_id: str, download: bool = False):
    """Export the latest proposal version as HTML (inline by default)."""
    try:
        latest = await get_latest_proposal_version_db(proposal_id)
        if not latest:
            raise HTTPException(status_code=404, detail="No versions found for this proposal")
        proposal = await get_proposal_db(proposal_id)
        title = proposal.get("title") if isinstance(proposal, dict) else f"Proposal {proposal_id}"
        html = _build_proposal_html(latest.get("sections"), title=title or f"Proposal {proposal_id}")
        headers = {}
        if download:
            headers["Content-Disposition"] = f"attachment; filename=proposal_{proposal_id}.html"
        return HTMLResponse(content=html, headers=headers)
    except HTTPException:
        raise
    except Exception as e:
        logger.exception("Export HTML failed: %s", e)
        raise HTTPException(status_code=500, detail="Failed to export proposal HTML")


@app.get("/proposals/{proposal_id}/export/docx")
async def api_export_proposal_docx(proposal_id: str, download: bool = True):
    """Export the latest proposal version as a DOCX file. Falls back to HTML if python-docx is unavailable."""
    try:
        latest = await get_latest_proposal_version_db(proposal_id)
        if not latest:
            raise HTTPException(status_code=404, detail="No versions found for this proposal")
        proposal = await get_proposal_db(proposal_id)
        title = proposal.get("title") if isinstance(proposal, dict) else f"Proposal {proposal_id}"

        try:
            from docx import Document  # type: ignore
        except Exception:
            # Fallback: return HTML download
            html = _build_proposal_html(latest.get("sections"), title=title or f"Proposal {proposal_id}")
            headers = {"Content-Disposition": f"attachment; filename=proposal_{proposal_id}.html"} if download else {}
            return HTMLResponse(content=html, headers=headers)

        doc = Document()
        doc.add_heading(title or f"Proposal {proposal_id}", level=1)
        for s in (latest.get("sections") or []):
            stitle = (s.get("title") or "Section") if isinstance(s, dict) else "Section"
            scontent = (s.get("content") or "") if isinstance(s, dict) else ""
            doc.add_heading(stitle, level=2)
            if scontent:
                # Strip basic HTML tags for docx body; naive fallback
                try:
                    from bs4 import BeautifulSoup  # optional, best-effort
                    text = BeautifulSoup(scontent, 'html.parser').get_text("\n")
                except Exception:
                    import re as _re
                    text = _re.sub(r"<[^>]+>", "", scontent)
                for para in (text or "").split("\n\n"):
                    if para.strip():
                        doc.add_paragraph(para.strip())

        buff = io.BytesIO()
        doc.save(buff)
        buff.seek(0)
        filename = f"proposal_{proposal_id}.docx"
        headers = {"Content-Disposition": f"attachment; filename={filename}"} if download else {}
        return StreamingResponse(buff, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", headers=headers)
    except HTTPException:
        raise
    except Exception as e:
        logger.exception("Export DOCX failed: %s", e)
        raise HTTPException(status_code=500, detail="Failed to export proposal DOCX")

@app.post("/proposals/{proposal_id}/generate/stream")
async def proposal_generate_stream(proposal_id: str, request: Request):
    """Stream generation for a single proposal section with retrieval SSE events.

    Body JSON fields:
      - section_title: title for the section
      - section_instructions: optional guidance
      - metadata: { contextMode, selectedCollections, selectedDocuments, selectedChunks, force_guided? }
      - search_type: 'hybrid' | 'vector' | 'graph' (optional)
    """
    try:
        body = await request.json()
        section_title = (body.get("section_title") or "Section").strip()
        section_key = (body.get("section_key") or None)
        section_instructions = (body.get("section_instructions") or "").strip()
        meta = body.get("metadata") or {}

        # Build search preferences from metadata
        prefs: Dict[str, Any] = {}
        context_mode = (meta.get("contextMode") or "all").lower()
        if context_mode == "collections":
            if isinstance(meta.get("selectedCollections"), list):
                prefs["collection_ids"] = meta.get("selectedCollections")
        if context_mode == "documents":
            if isinstance(meta.get("selectedDocuments"), list):
                prefs["document_ids"] = meta.get("selectedDocuments")
        if isinstance(meta.get("selectedChunks"), list):
            prefs["chunk_ids"] = meta.get("selectedChunks")

        # Create a transient session id for this stream
        session_id = str(uuid.uuid4())

        async def generate_stream():
            full_response = ""
            retrieval_queue = None
            guided_task = None
            try:
                # Register retrieval listener
                try:
                    retrieval_queue = register_retrieval_listener(session_id)
                except Exception:
                    retrieval_queue = None

                # Optionally run EnhancedRetriever in background to emit retrieval events
                force_guided = True if meta.get("force_guided", True) else False
                if force_guided:
                    async def _run_enhanced():
                        try:
                            retriever = EnhancedRetriever()
                            config: Dict[str, Any] = {
                                "use_graph": True,
                                "use_vector": True,
                                "use_query_expansion": True,
                                "vector_limit": 100,
                            }
                            # Apply scoping preferences
                            if prefs.get("collection_ids"):
                                config["collection_ids"] = prefs["collection_ids"]
                            if prefs.get("document_ids"):
                                config["document_ids"] = prefs["document_ids"]
                            if prefs.get("chunk_ids"):
                                config["chunk_ids"] = prefs["chunk_ids"]
                            await retriever.retrieve(
                                query=f"Generate section: {section_title}. {section_instructions}",
                                session_id=session_id,
                                config=config,
                            )
                        except Exception as e:
                            logger.warning(f"Enhanced retrieval failed in proposal stream: {e}")
                    try:
                        guided_task = asyncio.create_task(_run_enhanced())
                    except Exception:
                        guided_task = None

                # Build a generation prompt for the section
                # Include style guidance if the user uploaded an example proposal
                style_hint = (PROPOSAL_STYLE_HINTS.get(proposal_id, {}) or {}).get("style_prompt")
                phrase_bank = (PROPOSAL_STYLE_HINTS.get(proposal_id, {}) or {}).get("phrase_bank")
                formatting = (PROPOSAL_STYLE_HINTS.get(proposal_id, {}) or {}).get("formatting") or {}
                structure_analysis = (PROPOSAL_STYLE_HINTS.get(proposal_id, {}) or {}).get("structure_analysis", {})
                draft_text = PROPOSAL_DRAFT_TEXTS.get(proposal_id, "")

                # Retrieve top-rated examples for this section, if any
                top_examples: List[Dict[str, Any]] = []
                try:
                    if section_key:
                        top_examples = await get_top_rated_section_examples_db(section_key, limit=3)
                except Exception as e:
                    logger.debug(f"Fetching top-rated examples failed (non-fatal): {e}")

                # Prefetch scoped vector context snippets to ground the draft
                context_snippets: List[str] = []
                try:
                    vec_input = VectorSearchInput(
                        query=f"{section_title} {section_instructions}".strip() or section_title,
                        limit=12,
                        collection_ids=prefs.get("collection_ids"),
                        document_ids=prefs.get("document_ids"),
                        chunk_ids=prefs.get("chunk_ids"),
                    )
                    top_chunks = await vector_search_tool(vec_input)
                    for idx, ch in enumerate(top_chunks[:12]):
                        try:
                            citation = f"[{idx+1}] {ch.document_title or ''} — {ch.content.strip()}"
                            # Keep snippets reasonably short
                            if len(citation) > 800:
                                citation = citation[:800] + "…"
                            context_snippets.append(citation)
                        except Exception:
                            pass
                except Exception as e:
                    logger.debug(f"Prefetch vector context failed (non-fatal): {e}")

                drafting_policy = (
                    "You are in Proposal Drafting Mode. You are explicitly permitted to draft "
                    "professional proposal content. Do not refuse to draft due to 'creative content' "
                    "limitations. When specific facts are unavailable, write a polished, generic draft "
                    "and mark placeholders like <Client Name> or <Project Timeline>. Always prefer to ground "
                    "claims in the provided context snippets when relevant, and maintain a confident, helpful tone."
                )

                # Build optional formatting guide derived from example
                fmt_bits: List[str] = []
                bullet_marker = formatting.get("bullet_marker")
                if bullet_marker:
                    fmt_bits.append(f"Use this exact bullet character for lists: '{bullet_marker}' followed by a space.")
                if formatting.get("has_task_pattern"):
                    fmt_bits.append("For 'Scope of Work' sections, structure subsections as 'Task N: Title'.")
                if formatting.get("has_aoc_pattern"):
                    fmt_bits.append("When listing Areas of Concern, prefix headings as 'AOC-<number> <Title>' on their own lines.")
                if formatting.get("has_re_line") or formatting.get("has_salutation"):
                    fmt_bits.append("For cover letter style sections, begin with 'RE: <Subject>' and 'Dear <Recipient Name>:' on separate lines.")
                if fmt_bits:
                    formatting_guide = "Formatting Guide (mirror example layout):\n- " + "\n- ".join(fmt_bits)
                else:
                    formatting_guide = None

                # Section-specific guide to strongly enforce structure
                section_specific_guide = None
                try:
                    st_lower = (section_title or "").lower()
                    if "scope" in st_lower and "work" in st_lower:
                        bm = formatting.get("bullet_marker") or "-"
                        section_specific_guide = (
                            "For this section, structure subsections exactly as 'Task <number>: <Title>'. "
                            f"When listing items, use the bullet marker '{bm} ' at the start of each list line."
                        )
                    elif any(k in st_lower for k in ["regulatory", "compliance", "aoc"]):
                        bm = formatting.get("bullet_marker") or "-"
                        section_specific_guide = (
                            "When listing Areas of Concern, prefix headings as 'AOC-<number> <Title>' and use brief explanatory paragraphs. "
                            f"Use the bullet marker '{bm} ' for sublists where appropriate."
                        )
                    elif any(k in st_lower for k in ["cover", "letter"]):
                        section_specific_guide = (
                            "Begin with a header line 'RE: <Subject>' followed by a salutation 'Dear <Recipient Name>:' on the next line. "
                            "Write in formal letter style."
                        )
                except Exception:
                    section_specific_guide = None

                prompt_parts = [
                    drafting_policy,
                    f"Draft the '{section_title}' section of a professional proposal.",
                ]
                if section_instructions:
                    prompt_parts.append(f"Follow these instructions: {section_instructions}")
                if style_hint:
                    prompt_parts.append(f"Style guidance to mirror:\n{style_hint}")

                # Enhanced: Include Dolphin structure analysis insights
                if structure_analysis:
                    structure_guidance = []
                    if structure_analysis.get("has_tables"):
                        table_count = structure_analysis.get("table_count", 0)
                        structure_guidance.append(f"Document contains {table_count} table(s) - use tabular format when presenting structured data")
                    if structure_analysis.get("has_formulas"):
                        formula_count = structure_analysis.get("formula_count", 0)
                        structure_guidance.append(f"Document contains {formula_count} formula(s) - include mathematical notation where appropriate")

                    section_types = structure_analysis.get("section_types", {})
                    if section_types:
                        type_summary = ", ".join([f"{count} {stype}" for stype, count in section_types.items()])
                        structure_guidance.append(f"Document structure includes: {type_summary}")

                    if structure_guidance:
                        prompt_parts.append("Document structure insights from enhanced parsing:\n- " + "\n- ".join(structure_guidance))

                if formatting_guide:
                    prompt_parts.append(formatting_guide)
                if top_examples:
                    # Include brief exemplars distilled from feedback
                    ex_lines: List[str] = []
                    for i, ex in enumerate(top_examples[:3], start=1):
                        content = (ex.get("content") or "").strip()
                        title = (ex.get("title") or "").strip() or f"Example {i}"
                        snippet = content[:800] + ("…" if len(content) > 800 else "")
                        ex_lines.append(f"[{i}] {title}:\n{snippet}")
                    prompt_parts.append(
                        "Successful patterns from previously upvoted sections (use tone/structure appropriately, do not copy verbatim):\n" + "\n\n".join(ex_lines)
                    )
                if section_specific_guide:
                    prompt_parts.append(section_specific_guide)
                if phrase_bank and isinstance(phrase_bank, list) and phrase_bank:
                    prompt_parts.append(
                        "Preferred terminology (optional, from example): " + ", ".join(phrase_bank[:15])
                    )
                if context_snippets:
                    prompt_parts.append(
                        "Context snippets (cite with [n] when you use a fact):\n" + "\n\n".join(context_snippets)
                    )
                if draft_text:
                    # Trim draft content to avoid exceeding context; keep a generous but bounded portion
                    trimmed_draft = draft_text.strip()
                    if len(trimmed_draft) > 6000:
                        trimmed_draft = trimmed_draft[:6000] + "…"
                    prompt_parts.append(
                        "Existing draft content to adapt and improve (use as source material where relevant):\n" + trimmed_draft
                    )
                prompt_parts.append(
                    "Requirements:\n- Maintain consistent branding.\n- Use clear, concise paragraphs.\n- Include numbered citations [1], [2] when you rely on a context snippet.\n- If context is thin, still produce a useful draft with sensible placeholders."
                )
                prompt = "\n\n".join(prompt_parts)

                # Execute using rag_agent with scoping deps
                deps = AgentDependencies(
                    session_id=session_id,
                    user_id=request.headers.get("x-user-id"),
                    search_preferences=prefs,
                )

                async with rag_agent.iter(prompt, deps=deps) as run:
                    async for node in run:
                        if rag_agent.is_model_request_node(node):
                            async with node.stream(run.ctx) as request_stream:
                                from pydantic_ai.messages import PartStartEvent, PartDeltaEvent, TextPartDelta
                                async for event in request_stream:
                                    if isinstance(event, PartStartEvent) and event.part.part_kind == 'text':
                                        delta = event.part.content
                                        full_response += delta
                                        yield f"data: {json.dumps({'type': 'text', 'content': delta})}\n\n"
                                        # Drain retrieval events
                                        if retrieval_queue is not None:
                                            try:
                                                while True:
                                                    ev = retrieval_queue.get_nowait()
                                                    yield f"data: {json.dumps({'type': 'retrieval', 'session_id': session_id, 'data': ev})}\n\n"
                                            except asyncio.QueueEmpty:
            
                                                pass
                                    elif isinstance(event, PartDeltaEvent) and isinstance(event.delta, TextPartDelta):
                                        delta = event.delta.content_delta
                                        full_response += delta
                                        yield f"data: {json.dumps({'type': 'text', 'content': delta})}\n\n"
                                        if retrieval_queue is not None:
                                            try:
                                                while True:
                                                    ev = retrieval_queue.get_nowait()
                                                    yield f"data: {json.dumps({'type': 'retrieval', 'session_id': session_id, 'data': ev})}\n\n"
                                            except asyncio.QueueEmpty:
                                                pass

                    # Give guided task a moment to finish
                    if guided_task is not None:
                        try:
                            await asyncio.wait_for(guided_task, timeout=1.0)
                        except Exception:
                            pass
                    # Final drain
                    if retrieval_queue is not None:
                        try:
                            while True:
                                ev = retrieval_queue.get_nowait()
                                yield f"data: {json.dumps({'type': 'retrieval', 'session_id': session_id, 'data': ev})}\n\n"
                        except asyncio.QueueEmpty:
                            pass

                # End event
                yield f"data: {json.dumps({'type': 'end'})}\n\n"
            except Exception as e:
                yield f"data: {json.dumps({'type': 'error', 'error': str(e)})}\n\n"
                yield f"data: {json.dumps({'type': 'end'})}\n\n"
            finally:
                try:
                    if retrieval_queue is not None:
                        unregister_retrieval_listener(session_id, retrieval_queue)
                except Exception:
                    pass

        return StreamingResponse(
            generate_stream(),
            media_type="text/event-stream",
            headers={
                "Cache-Control": "no-cache",
                "Connection": "keep-alive",
                "Content-Type": "text/event-stream",
            },
        )
    except Exception as e:
        logger.exception("Proposal generate stream failed: %s", e)
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/proposals/{proposal_id}/aocs/discover")
async def proposal_aoc_discover(proposal_id: str, request: Request):
    """Discover AOC codes and optional titles from proposal-scoped context.

    Body JSON fields (optional):
      - contextMode: 'all' | 'collections' | 'documents'
      - selectedCollections: [ids]
      - selectedDocuments: [ids]
      - limit: number of chunks to scan (default 150)
    """
    try:
        body = await request.json()
    except Exception:
        body = {}

    context_mode = (body.get("contextMode") or "all").lower()
    collection_ids = body.get("selectedCollections") or []
    document_ids = body.get("selectedDocuments") or []
    limit = int(body.get("limit") or 150)

    # Build search preferences
    prefs: Dict[str, Any] = {}
    if context_mode == "collections" and isinstance(collection_ids, list):
        prefs["collection_ids"] = collection_ids
    if context_mode == "documents" and isinstance(document_ids, list):
        prefs["document_ids"] = document_ids

    # Vector search with queries that likely surface AOC mentions
    queries = [
        "AOC-",
        "Areas of Concern (AOC)",
        "AOC code",
    ]
    seen_codes = set()
    aocs: List[Dict[str, str]] = []

    code_re = re.compile(r"AOC[-\s]?([0-9]+[A-Za-z]?)", re.IGNORECASE)

    async def scan_query(q: str):
        try:
            vi = VectorSearchInput(
                query=q,
                limit=limit,
                collection_ids=prefs.get("collection_ids"),
                document_ids=prefs.get("document_ids"),
            )
            hits = await vector_search_tool(vi)
            for h in hits or []:
                text = (getattr(h, "content", None) or "")
                for m in code_re.finditer(text):
                    code = f"AOC-{m.group(1).upper()}"
                    if code in seen_codes:
                        continue
                    # Try to grab a short title after the code in the same sentence
                    tail = text[m.end(): m.end()+140]
                    title = None
                    # look for dash/colon patterns
                    dash_idx = tail.find("-")
                    colon_idx = tail.find(":")
                    sep_idx = dash_idx if (dash_idx != -1 and (colon_idx == -1 or dash_idx < colon_idx)) else colon_idx
                    if sep_idx != -1:
                        cand = tail[sep_idx+1:].strip().split(".\n")[0].split(". ")[0]
                        cand = cand[:80].strip()
                        title = cand
                    aocs.append({"code": code, **({"title": title} if title else {})})
                    seen_codes.add(code)
                    if len(seen_codes) >= 100:
                        return
        except Exception as e:
            logger.debug(f"AOC discover scan failed for query '{q}': {e}")

    # Run scans sequentially to respect DB pool limits
    for q in queries:
        await scan_query(q)
        if len(seen_codes) >= 100:
            break

    # Also scan uploaded draft text (if any) for inline AOC headings
    try:
        draft_text = (PROPOSAL_DRAFT_TEXTS.get(proposal_id) or "").strip()
        if draft_text:
            def _clean_title(raw: str) -> str:
                t = (raw or "").strip()
                if not t:
                    return ""
                # Remove dot leaders (contiguous or spaced) and bullet-like dots
                t = re.sub(r"(?:\s*[\.\u2024\u2027\u2219\u00B7•\u2022]\s*){3,}", "", t)
                # Remove stray trailing "- Back-" fragments
                t = re.sub(r"\s*[-–—]\s*Back\s*[-–—]?\s*$", "", t, flags=re.IGNORECASE)
                # Collapse multiple spaces
                t = re.sub(r"\s{2,}", " ", t).strip()
                # Trim trailing and leading punctuation
                t = re.sub(r"[\s\.\-–—:;,]+$", "", t).strip()
                t = re.sub(r"^[\s\.\-–—:;,]+", "", t).strip()
                return t[:80]

            for m in code_re.finditer(draft_text):
                code = f"AOC-{m.group(1).upper()}"
                if code in seen_codes:
                    continue
                # Heuristic title extraction from the same line
                line_start = draft_text.rfind('\n', 0, m.start()) + 1
                line_end_n = draft_text.find('\n', m.end())
                line_end = line_end_n if line_end_n != -1 else len(draft_text)
                line = draft_text[line_start:line_end]
                title = None
                # Take the remainder of the line after the code token.
                # If there is an immediate separator (dash/en dash/em dash/colon) right after the code, strip only that leading separator.
                after = line[m.end()-line_start:]
                after2 = re.sub(r"^\s*[-–—:]\s*", "", after)
                cand = after2.strip()
                # Split off page leaders / repeated dots
                cand = re.split(r"\s?\.{2,}\s?", cand)[0]
                # And line breaks after a period
                cand = cand.split('.\n')[0]
                cand = _clean_title(cand)
                title = cand or None
                aocs.append({"code": code, **({"title": title} if title else {})})
                seen_codes.add(code)
    except Exception as e:
        logger.debug(f"Draft AOC scan failed: {e}")

    # If nothing found, try looking up proposal documents and lightly seed
    if not aocs:
        try:
            docs = await list_proposal_documents_db(proposal_id, source_type=None)
            ids = [d.get("id") or d.get("document_id") for d in (docs or []) if (d.get("id") or d.get("document_id"))]
            if ids:
                vi = VectorSearchInput(query="AOC-", limit=limit, document_ids=ids)
                hits = await vector_search_tool(vi)
                for h in hits or []:
                    text = (getattr(h, "content", None) or "")
                    for m in code_re.finditer(text):
                        code = f"AOC-{m.group(1).upper()}"
                        if code in seen_codes:
                            continue
                        aocs.append({"code": code})
                        seen_codes.add(code)
        except Exception as e:
            logger.debug(f"Fallback proposal doc scan failed: {e}")

    # Return deduped list
    return {"aocs": aocs}


@app.post("/pricing/parse")
async def pricing_parse(file: UploadFile = File(...)):
    """Parse CSV pricing file into items array."""
    try:
        # Only CSV support for now
        filename = file.filename or "pricing.csv"
        if not filename.lower().endswith('.csv'):
            raise HTTPException(status_code=415, detail="Only CSV is supported at the moment")
        # Read CSV text
        content = (await file.read()).decode('utf-8', errors='ignore')
        reader = csv.DictReader(content.splitlines())
        items = []
        for row in reader:
            service = (row.get('service') or row.get('Service') or '').strip()
            qty = row.get('quantity') or row.get('qty') or row.get('Quantity') or '1'
            unit = row.get('unit_price') or row.get('price') or row.get('Unit Price') or '0'
            desc = row.get('description') or row.get('Description') or ''
            cur = row.get('currency_symbol') or row.get('Currency') or '$'
            try:
                quantity = float(str(qty).strip() or '1')
            except Exception:
                quantity = 1.0
            try:
                unit_price = float(str(unit).strip() or '0')
            except Exception:
                unit_price = 0.0
            items.append({
                "service": service,
                "unit_price": unit_price,
                "quantity": quantity,
                "description": desc,
                "currency_symbol": cur or '$'
            })
        return {"items": items}
    except HTTPException:
        raise
    except Exception as e:
        logger.exception("Failed to parse pricing file: %s", e)
        raise HTTPException(status_code=500, detail="Failed to parse pricing file")


@app.post("/pricing/render")
async def pricing_render(request: Request):
    """Render a pricing table to HTML and totals."""
    try:
        body = await request.json()
        items = body.get("items") or []
        tax_rate = float(body.get("tax_rate_percent") or 0)
        discount = float(body.get("discount_amount") or 0)
        subtotal = 0.0
        currency = '$'
        for it in items:
            currency = it.get('currency_symbol') or currency
            try:
                subtotal += float(it.get('unit_price') or 0) * float(it.get('quantity') or 0)
            except Exception:
                pass
        tax = subtotal * (tax_rate / 100.0)
        total = max(0.0, subtotal + tax - discount)

        # Create simple HTML table
        rows = []
        for it in items:
            rows.append(
                f"<tr><td>{it.get('service','')}</td><td style='text-align:right'>{it.get('quantity',0)}</td>"
                f"<td style='text-align:right'>{currency}{float(it.get('unit_price') or 0):.2f}</td>"
                f"<td>{(it.get('description') or '')}</td></tr>"
            )
        totals_html = (
            f"<div style='margin-top:8px;text-align:right'>"
            f"<div>Subtotal: {currency}{subtotal:.2f}</div>"
            f"<div>Tax ({tax_rate:.2f}%): {currency}{tax:.2f}</div>"
            f"<div>Discount: -{currency}{discount:.2f}</div>"
            f"<div><strong>Total: {currency}{total:.2f}</strong></div>"
            f"</div>"
        )
        html = (
            "<table style='width:100%;border-collapse:collapse'>"
            "<thead><tr><th style='text-align:left'>Service</th><th style='text-align:right'>Qty</th><th style='text-align:right'>Unit Price</th><th style='text-align:left'>Description</th></tr></thead>"
            "<tbody>" + "".join(rows) + "</tbody>" + "</table>" + totals_html
        )
        return {"html": html, "totals": {"subtotal": subtotal, "tax": tax, "discount": discount, "total": total}}
    except Exception as e:
        logger.exception("Failed to render pricing: %s", e)
        raise HTTPException(status_code=500, detail="Failed to render pricing")

# -------------------------
# Phase 1: Incremental Update Endpoints (metadata-only, no embeddings)
# -------------------------

@app.patch("/documents/{document_id}/metadata")
async def patch_document_metadata(document_id: str, req: DocumentMetadataUpdateRequest):
    """Merge metadata into a document without re-embedding."""
    try:
        if not req or not isinstance(req.metadata, dict) or not req.metadata:
            raise HTTPException(status_code=400, detail="metadata is required and must be a non-empty object")
        updated = await update_document_metadata_only(document_id, req.metadata)
        if not updated:
            raise HTTPException(status_code=404, detail="Document not found")
        return {"document_id": updated["id"], "metadata": updated["metadata"], "updated_at": updated["updated_at"]}
    except HTTPException:
        raise
    except Exception as e:
        logger.exception("Failed to update document metadata: %s", e)
        raise HTTPException(status_code=500, detail="Failed to update document metadata")


@app.patch("/chunks/metadata")
async def patch_chunks_metadata(req: ChunkMetadataBatchUpdateRequest):
    """Bulk-merge metadata for multiple chunks without re-embedding."""
    try:
        if not req or not req.chunk_ids or not isinstance(req.metadata, dict) or not req.metadata:
            raise HTTPException(status_code=400, detail="chunk_ids and non-empty metadata are required")
        updated_count = await update_chunk_metadata_batch(req.chunk_ids, req.metadata)
        return {"updated": updated_count}
    except HTTPException:
        raise
    except Exception as e:
        logger.exception("Failed to update chunk metadata batch: %s", e)
        raise HTTPException(status_code=500, detail="Failed to update chunk metadata")


@app.post("/documents/{document_id}/tags")
async def post_document_tags(document_id: str, req: AddTagsRequest):
    """Add tags to a document (idempotent)."""
    try:
        if not req or not req.tags:
            raise HTTPException(status_code=400, detail="tags array is required")
        added = await add_document_tags(document_id, req.tags)
        return {"document_id": document_id, "added": added}
    except HTTPException:
        raise
    except Exception as e:
        logger.exception("Failed to add document tags: %s", e)
        raise HTTPException(status_code=500, detail="Failed to add document tags")


@app.post("/documents/{document_id}/classification")
async def post_document_classification(document_id: str, req: UpdateClassificationRequest):
    """Update document domain classification and optional chunk categories without re-embedding."""
    try:
        total_chunk_updates = 0
        # Update document-level classification in metadata if provided
        doc_meta: Dict[str, Any] = {}
        if req and req.domain:
            doc_meta["domain"] = req.domain
        if req and req.domain_confidence is not None:
            doc_meta["domain_confidence"] = req.domain_confidence
        if doc_meta:
            updated = await update_document_metadata_only(document_id, doc_meta)
            if not updated:
                raise HTTPException(status_code=404, detail="Document not found")

        # Update chunk categories in batch per group
        if req and req.chunk_category_updates:
            for update in req.chunk_category_updates:
                if update.chunk_ids and update.category:
                    total_chunk_updates += await update_chunk_metadata_batch(update.chunk_ids, {"category": update.category})

        return {
            "document_id": document_id,
            "chunk_category_updates": total_chunk_updates,
        }
    except HTTPException:
        raise
    except Exception as e:
        logger.exception("Failed to update classification: %s", e)
        raise HTTPException(status_code=500, detail="Failed to update classification")


@app.put("/documents/{document_id}/collections")
async def put_document_collections(document_id: str, req: UpdateCollectionsRequest, request: Request):
    """Set the exact set of collections for a document (add/remove memberships)."""
    try:
        if not req or not isinstance(req.collection_ids, list):
            raise HTTPException(status_code=400, detail="collection_ids array is required")
        actor = request.headers.get("x-user-id")
        result = await update_document_collections_db(document_id, req.collection_ids, added_by=actor)
        return {"document_id": document_id, **result}
    except HTTPException:
        raise
    except Exception as e:
        logger.exception("Failed to update document collections: %s", e)
        raise HTTPException(status_code=500, detail="Failed to update document collections")


@app.delete("/collections/{collection_id}/documents/{document_id}")
async def api_remove_document_from_collection(collection_id: str, document_id: str):
    try:
        ok = await remove_document_from_collection_db(collection_id, document_id)
        if not ok:
            raise HTTPException(status_code=404, detail="Not found")
        return {"removed": True}
    except HTTPException:
        raise
    except Exception as e:
        logger.exception("Failed to remove document from collection: %s", e)
        raise HTTPException(status_code=500, detail="Failed to remove document from collection")
@app.get("/health", response_model=HealthStatus)
async def health_check():
    """Health check endpoint."""
    try:
        # Test database connections
        db_status = await test_connection()
        graph_status = await test_graph_connection()
        
        # Determine overall status
        if db_status and graph_status:
            status = "healthy"
        elif db_status or graph_status:
            status = "degraded"
        else:
            status = "unhealthy"
        
        return HealthStatus(
            status=status,
            database=db_status,
            graph_database=graph_status,
            llm_connection=True,  # Assume OK if we can respond
            version="0.1.0",
            timestamp=datetime.now()
        )
        
    except Exception as e:
        logger.error(f"Health check failed: {e}")
        raise HTTPException(status_code=500, detail="Health check failed")


@app.post("/chat", response_model=ChatResponse)
async def chat(request: ChatRequest):
    """Non-streaming chat endpoint."""
    try:
        # Get or create session
        session_id = await get_or_create_session(request)
        
        # Build search preferences from metadata
        prefs: Dict[str, Any] = {}
        meta = request.metadata or {}
        # Support multiple key styles from client
        if isinstance(meta.get("collection_ids"), list):
            prefs["collection_ids"] = meta.get("collection_ids")
        if isinstance(meta.get("selectedCollections"), list):
            prefs["collection_ids"] = meta.get("selectedCollections")
        if isinstance(meta.get("document_ids"), list):
            prefs["document_ids"] = meta.get("document_ids")
        if isinstance(meta.get("selectedDocuments"), list):
            prefs["document_ids"] = meta.get("selectedDocuments")

        # Execute agent
        response, tools_used, sources = await execute_agent(
            message=request.message,
            session_id=session_id,
            user_id=request.user_id,
            search_preferences=prefs,
        )
        
        # Build response metadata and attach a live graph payload when appropriate
        response_metadata: Dict[str, Any] = {"search_type": str(request.search_type)}

        graph_included = False
        try:
            wants_graph = "graph" in (request.message or "").lower()
            used_graph_tool = any(
                getattr(t, "tool_name", "").lower().find("graph") != -1 for t in tools_used
            )
            if wants_graph or used_graph_tool:
                live_graph = await build_live_graph_payload(request.message, tools_used)
                if live_graph and (live_graph.get("nodes") or live_graph.get("edges")):
                    response_metadata["graph"] = live_graph
                    graph_included = True
        except Exception:
            # Fallback: don't block the response if anything goes wrong building the graph payload
            pass

        # If we attached a graph, ensure the assistant message acknowledges it and remove contradictory phrasing
        if graph_included:
            try:
                safe_response = response or ""
                lower_resp = safe_response.lower()
                negative_phrases = [
                    "unable to generate a visual graph",
                    "cannot generate a visual graph",
                    "can not generate a visual graph",
                    "unable to generate a graph",
                    "cannot generate a graph",
                    "can not generate a graph",
                    "i am unable to generate a graph",
                    "i cannot generate a graph",
                    "i'm unable to generate a graph",
                ]
                for phrase in negative_phrases:
                    if phrase in lower_resp:
                        # Replace the first occurrence with a positive confirmation
                        idx = lower_resp.find(phrase)
                        safe_response = (
                            safe_response[:idx]
                            + "I've generated an interactive knowledge graph below."
                            + safe_response[idx + len(phrase):]
                        )
                        lower_resp = safe_response.lower()

                if "knowledge graph" not in lower_resp:
                    safe_response = (
                        "I've generated an interactive knowledge graph below.\n\n"
                        + safe_response
                    )
                response = safe_response
            except Exception:
                # Don't fail the request due to message post-processing
                pass

        return ChatResponse(
            message=response,
            session_id=session_id,
            sources=sources,
            tools_used=tools_used,
            metadata=response_metadata
        )
        
    except Exception as e:
        logger.error(f"Chat endpoint failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/chat/stream")
async def chat_stream(request: ChatRequest):
    """Compatibility route for /chat/stream that forwards to the existing chat_stream_legacy."""
    return await chat_stream_legacy(request)


@app.post("/chat/stream/legacy")  # Renamed to avoid conflict with enhanced endpoint
async def chat_stream_legacy(request: ChatRequest):
    """Streaming chat endpoint using Server-Sent Events."""
    try:
        # Get or create session
        session_id = await get_or_create_session(request)
        
        async def generate_stream():
            """Generate streaming response using agent.iter() pattern."""
            try:
                yield f"data: {json.dumps({'type': 'session', 'session_id': session_id})}\n\n"
                
                # Create dependencies
                # Build search preferences from metadata
                prefs: Dict[str, Any] = {}
                meta = request.metadata or {}
                if isinstance(meta.get("collection_ids"), list):
                    prefs["collection_ids"] = meta.get("collection_ids")
                if isinstance(meta.get("selectedCollections"), list):
                    prefs["collection_ids"] = meta.get("selectedCollections")
                if isinstance(meta.get("document_ids"), list):
                    prefs["document_ids"] = meta.get("document_ids")
                if isinstance(meta.get("selectedDocuments"), list):
                    prefs["document_ids"] = meta.get("selectedDocuments")

                deps = AgentDependencies(
                    session_id=session_id,
                    user_id=request.user_id,
                    search_preferences=prefs,
                )
                
                # Get conversation context
                context = await get_conversation_context(session_id)
                
                # Build input with context
                full_prompt = request.message
                if context:
                    context_str = "\n".join([
                        f"{msg['role']}: {msg['content']}"
                        for msg in context[-6:]
                    ])
                    full_prompt = f"Previous conversation:\n{context_str}\n\nCurrent question: {request.message}"
                
                # Save user message immediately (best-effort)
                try:
                    await add_message(
                        session_id=session_id,
                        role="user",
                        content=request.message,
                        metadata={"user_id": request.user_id}
                    )
                except Exception as e:
                    logger.warning("Failed to persist user message; continuing stream: %s", e)
                
                full_response = ""
                # Register live retrieval events listener for this session (graceful fallback)
                retrieval_queue = None
                guided_task = None
                try:
                    retrieval_queue = register_retrieval_listener(session_id)
                    logger.info(f"Registered retrieval listener for session {session_id}, queue={retrieval_queue}")
                except Exception as e:
                    logger.error(f"Failed to register retrieval listener: {e}")
                    retrieval_queue = None
                
                # Determine if mock streaming mode is enabled
                use_mock = False
                try:
                    use_mock = bool(
                        (getattr(request, "metadata", {}) or {}).get("mock_stream")
                        or (getattr(request, "metadata", {}) or {}).get("mock")
                        or os.getenv("MOCK_STREAM") == "1"
                    )
                except Exception:
                    use_mock = False

                # Optionally force Enhanced Graph → Vector retrieval to run and emit retrieval_step events
                force_guided = False
                try:
                    force_guided = bool(
                        ((getattr(request, "metadata", {}) or {}).get("force_guided"))
                        or os.getenv("FORCE_GUIDED_RETRIEVAL") == "1"
                    )
                except Exception:
                    force_guided = os.getenv("FORCE_GUIDED_RETRIEVAL") == "1"

                if force_guided and not use_mock:
                    async def _run_enhanced_retrieval():
                        try:
                            retriever = EnhancedRetriever()
                            config = {
                                "use_graph": True,
                                "use_vector": True,
                                "use_query_expansion": True,
                                "vector_limit": 100,  # Optimized for 1M context
                            }
                            # This will emit granular retrieval_step events via emit_retrieval_event
                            await retriever.retrieve(
                                query=request.message,
                                session_id=session_id,
                                config=config,
                            )
                        except Exception as e:
                            logger.warning(f"Forced guided retrieval failed: {e}")

                    try:
                        guided_task = asyncio.create_task(_run_enhanced_retrieval())
                    except Exception:
                        guided_task = None

                if use_mock:
                    # Emit mock staged retrieval events (guided_retrieval: graph -> vector)
                    async def _emit_mock_retrieval_events():
                        try:
                            # Orchestrator start
                            await emit_retrieval_event(session_id, {
                                "type": "retrieval",
                                "event": "start",
                                "tool": "guided_retrieval",
                                "args": {"query": request.message, "limit": 50}  # Increased for 1M context
                            })

                            # Graph stage
                            await emit_retrieval_event(session_id, {
                                "type": "retrieval",
                                "event": "start",
                                "tool": "guided_retrieval",
                                "stage": "graph"
                            })
                            await asyncio.sleep(0.05)
                            mock_graph = [
                                {"fact": "Company X acquired Startup Y in 2023", "uuid": "g-1"},
                                {"fact": "Startup Y builds vector databases", "uuid": "g-2"}
                            ]
                            await emit_retrieval_event(session_id, {
                                "type": "retrieval",
                                "event": "results",
                                "tool": "guided_retrieval",
                                "stage": "graph",
                                "results": mock_graph
                            })
                            await emit_retrieval_event(session_id, {
                                "type": "retrieval",
                                "event": "end",
                                "tool": "guided_retrieval",
                                "stage": "graph",
                                "count": len(mock_graph),
                                "elapsed_ms": 60
                            })

                            # Vector stage
                            await emit_retrieval_event(session_id, {
                                "type": "retrieval",
                                "event": "start",
                                "tool": "guided_retrieval",
                                "stage": "vector",
                                "args": {"limit": 30}  # Increased for 1M context
                            })
                            await asyncio.sleep(0.05)
                            mock_vector = [
                                {"content": "Mock result A", "score": 0.91, "document_title": "Doc A", "document_source": "mock", "chunk_id": "mock-a"},
                                {"content": "Mock result B", "score": 0.83, "document_title": "Doc B", "document_source": "mock", "chunk_id": "mock-b"}
                            ]
                            await emit_retrieval_event(session_id, {
                                "type": "retrieval",
                                "event": "results",
                                "tool": "guided_retrieval",
                                "stage": "vector",
                                "results": mock_vector
                            })
                            await emit_retrieval_event(session_id, {
                                "type": "retrieval",
                                "event": "end",
                                "tool": "guided_retrieval",
                                "stage": "vector",
                                "count": len(mock_vector),
                                "elapsed_ms": 70
                            })

                            # Orchestrator end
                            await emit_retrieval_event(session_id, {
                                "type": "retrieval",
                                "event": "end",
                                "tool": "guided_retrieval",
                                "count": len(mock_graph) + len(mock_vector),
                                "elapsed_ms": 140
                            })
                        except Exception:
                            # Never let mock emissions break streaming
                            pass

                    try:
                        asyncio.create_task(_emit_mock_retrieval_events())
                    except Exception:
                        pass

                    # Stream mock token chunks
                    for chunk in ["This ", "is ", "a ", "mock ", "stream ", "response."]:
                        await asyncio.sleep(0.05)
                        yield f"data: {json.dumps({'type': 'text', 'content': chunk})}\n\n"
                        full_response += chunk
                        # Drain any pending retrieval events and stream them
                        if retrieval_queue is not None:
                            try:
                                while True:
                                    ev = retrieval_queue.get_nowait()
                                    yield f"data: {json.dumps({'type': 'retrieval', 'session_id': session_id, 'data': ev})}\n\n"
                            except asyncio.QueueEmpty:
                                pass

                    # No tool calls when in mock mode
                    tools_used = []
                else:
                    # Stream using agent.iter() pattern
                    async with rag_agent.iter(full_prompt, deps=deps) as run:
                        async for node in run:
                            if rag_agent.is_model_request_node(node):
                                # Stream tokens from the model
                                async with node.stream(run.ctx) as request_stream:
                                    async for event in request_stream:
                                        from pydantic_ai.messages import PartStartEvent, PartDeltaEvent, TextPartDelta
                                        
                                        if isinstance(event, PartStartEvent) and event.part.part_kind == 'text':
                                            delta_content = event.part.content
                                            yield f"data: {json.dumps({'type': 'text', 'content': delta_content})}\n\n"
                                            full_response += delta_content
                                            
                                            # Drain any pending retrieval events and stream them
                                            if retrieval_queue is not None:
                                                try:
                                                    while True:
                                                        ev = retrieval_queue.get_nowait()
                                                        yield f"data: {json.dumps({'type': 'retrieval', 'session_id': session_id, 'data': ev})}\n\n"
                                                except asyncio.QueueEmpty:
                                                    pass
                                            
                                        elif isinstance(event, PartDeltaEvent) and isinstance(event.delta, TextPartDelta):
                                            delta_content = event.delta.content_delta
                                            yield f"data: {json.dumps({'type': 'text', 'content': delta_content})}\n\n"
                                            full_response += delta_content
                                            
                                            # Drain any pending retrieval events and stream them
                                            if retrieval_queue is not None:
                                                try:
                                                    while True:
                                                        ev = retrieval_queue.get_nowait()
                                                        yield f"data: {json.dumps({'type': 'retrieval', 'session_id': session_id, 'data': ev})}\n\n"
                                                except asyncio.QueueEmpty:
                                                    pass
                    
                    # Extract tools used from the final result
                    result = run.result
                    
                    # Fallback: if no token chunks were streamed but we have a final result,
                    # emit it as a single text event so the client gets an answer.
                    try:
                        final_text = getattr(result, "data", None)
                        if (not full_response.strip()) and isinstance(final_text, str) and final_text.strip():
                            yield f"data: {json.dumps({'type': 'text', 'content': final_text})}\n\n"
                            full_response += final_text
                    except Exception:
                        pass

                    tools_used = extract_tool_calls(result)

                    # If we launched guided retrieval, give it a brief moment to finish and flush events
                    if guided_task is not None:
                        try:
                            await asyncio.wait_for(guided_task, timeout=1.0)
                        except Exception:
                            pass

                    # Final drain of retrieval events to ensure 'complete' and 'summary' reach the client
                    if retrieval_queue is not None:
                        try:
                            while True:
                                ev = retrieval_queue.get_nowait()
                                yield f"data: {json.dumps({'type': 'retrieval', 'session_id': session_id, 'data': ev})}\n\n"
                        except asyncio.QueueEmpty:
                            pass
                
                # Final drain of retrieval events after model finished
                if retrieval_queue is not None:
                    try:
                        while True:
                            ev = retrieval_queue.get_nowait()
                            yield f"data: {json.dumps({'type': 'retrieval', 'session_id': session_id, 'data': ev})}\n\n"
                    except asyncio.QueueEmpty:
                        pass

                # Send tools used information
                if tools_used:
                    tools_data = [
                        {
                            "tool_name": tool.tool_name,
                            "args": tool.args,
                            "tool_call_id": tool.tool_call_id
                        }
                        for tool in tools_used
                    ]
                    yield f"data: {json.dumps({'type': 'tools', 'tools': tools_data})}\n\n"
                
                # Optionally build and stream a live knowledge graph payload
                graph_included = False
                live_graph = None
                try:
                    wants_graph = "graph" in (request.message or "").lower()
                    used_graph_tool = any(
                        getattr(t, "tool_name", "").lower().find("graph") != -1 for t in tools_used
                    )
                    if wants_graph or used_graph_tool:
                        live_graph = await build_live_graph_payload(request.message, tools_used)
                        if live_graph and (live_graph.get("nodes") or live_graph.get("edges")):
                            # Stream the graph payload as its own SSE event
                            yield f"data: {json.dumps({'type': 'graph', 'graph': live_graph})}\n\n"
                            graph_included = True
                except Exception:
                    # Never block streaming on graph construction issues
                    pass
                
                # If we attached a graph, adjust the assistant message to acknowledge it
                if graph_included:
                    try:
                        safe_response = full_response or ""
                        lower_resp = safe_response.lower()
                        negative_phrases = [
                            "unable to generate a visual graph",
                            "cannot generate a visual graph",
                            "can not generate a visual graph",
                            "unable to generate a graph",
                            "cannot generate a graph",
                            "can not generate a graph",
                            "i am unable to generate a graph",
                            "i cannot generate a graph",
                            "i'm unable to generate a graph",
                        ]
                        for phrase in negative_phrases:
                            if phrase in lower_resp:
                                idx = lower_resp.find(phrase)
                                safe_response = (
                                    safe_response[:idx]
                                    + "I've generated an interactive knowledge graph below."
                                    + safe_response[idx + len(phrase):]
                                )
                                lower_resp = safe_response.lower()
                        if "knowledge graph" not in lower_resp:
                            safe_response = (
                                "I've generated an interactive knowledge graph below.\n\n"
                                + safe_response
                            )
                        response_for_save = safe_response
                    except Exception:
                        # Do not fail the stream due to message post-processing
                        pass
                else:
                    response_for_save = full_response

                # Save assistant message after streaming completes (best-effort)
                try:
                    await add_message(
                        session_id=session_id,
                        role="assistant",
                        content=response_for_save,
                        metadata={"user_id": request.user_id}
                    )
                except Exception as e:
                    logger.warning("Failed to persist assistant message; continuing: %s", e)

                # Final end event
                yield f"data: {json.dumps({'type': 'end'})}\n\n"
            except Exception as e:
                # Stream error event and end
                yield f"data: {json.dumps({'type': 'error', 'error': str(e)})}\n\n"
                yield f"data: {json.dumps({'type': 'end'})}\n\n"
            finally:
                # Ensure we always unregister the retrieval listener
                try:
                    if retrieval_queue is not None:
                        unregister_retrieval_listener(session_id, retrieval_queue)
                except Exception:
                    pass
        
        return StreamingResponse(
            generate_stream(),
            media_type="text/event-stream",
            headers={
                "Cache-Control": "no-cache",
                "Connection": "keep-alive",
                "Content-Type": "text/event-stream"
            }
        )
        
    except Exception as e:
        logger.error(f"Streaming chat failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/search/vector")
async def search_vector(request: SearchRequest):
    """Vector search endpoint."""
    try:
        input_data = VectorSearchInput(
            query=request.query,
            limit=request.limit
        )
        
        start_time = datetime.now()
        results = await vector_search_tool(input_data)
        end_time = datetime.now()
        
        query_time = (end_time - start_time).total_seconds() * 1000
        
        return SearchResponse(
            results=results,
            total_results=len(results),
            search_type="vector",
            query_time_ms=query_time
        )
        
    except Exception as e:
        logger.error(f"Vector search failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/search/graph")
async def search_graph(request: SearchRequest):
    """Knowledge graph search endpoint."""
    try:
        # Use query understanding to make NL queries graph-friendly
        try:
            qp = QueryProcessor()
            processed = qp.process(request.query)
            effective_query = processed.graph_query or processed.cleaned or request.query
        except Exception:
            effective_query = request.query

        input_data = GraphSearchInput(
            query=effective_query
        )

        start_time = datetime.now()
        results = await graph_search_tool(input_data)

        # Hard fallback: if no results via tool, hit DB directly with broader websearch
        if not results:
            try:
                # First try websearch on the optimized query
                raw = await search_facts_websearch(effective_query, request.limit)
                # If still nothing and we altered the query, try the original as a last resort
                if not raw and effective_query != request.query:
                    raw = await search_facts_websearch(request.query, request.limit)
                fallback = [
                    GraphSearchResult(
                        fact=r.get("content", ""),
                        uuid=str(r.get("fact_id")),
                        valid_at=r.get("valid_at"),
                        invalid_at=r.get("invalid_at"),
                        source_node_uuid=str(r.get("node_id")) if r.get("node_id") else None
                    ) for r in raw
                ]
                results = fallback
            except Exception as fe:
                logger.warning(f"KG fallback (websearch) failed: {fe}")

        end_time = datetime.now()
        query_time = (end_time - start_time).total_seconds() * 1000

        return SearchResponse(
            graph_results=results,
            total_results=len(results),
            search_type="graph",
            query_time_ms=query_time
        )
        
    except Exception as e:
        logger.error(f"Graph search failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/graph/stats")
async def graph_stats():
    """Return live knowledge graph statistics (nodes, edges, facts)."""
    try:
        stats = await get_graph_statistics()
        # Ensure a predictable JSON shape
        return {
            "graph_initialized": bool(stats.get("graph_initialized", True)),
            "total_nodes": int(stats.get("total_nodes", 0) or 0),
            "total_edges": int(stats.get("total_edges", 0) or 0),
            "total_facts": int(stats.get("total_facts", 0) or 0),
            "nodes_by_type": stats.get("nodes_by_type", {})
        }
    except Exception as e:
        logger.error(f"Graph stats failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/search/hybrid")
async def search_hybrid(request: SearchRequest):
    """Hybrid search endpoint."""
    try:
        input_data = HybridSearchInput(
            query=request.query,
            limit=request.limit
        )
        
        start_time = datetime.now()
        results = await hybrid_search_tool(input_data)
        end_time = datetime.now()
        
        query_time = (end_time - start_time).total_seconds() * 1000
        
        return SearchResponse(
            results=results,
            total_results=len(results),
            search_type="hybrid",
            query_time_ms=query_time
        )
        
    except Exception as e:
        logger.error(f"Hybrid search failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/documents")
async def list_documents_endpoint(
    limit: int = 20,
    offset: int = 0
):
    """List documents endpoint."""
    try:
        input_data = DocumentListInput(limit=limit, offset=offset)
        documents = await list_documents_tool(input_data)
        
        # Get total count from database
        from .db_utils import db_pool
        async with db_pool.acquire() as conn:
            total_count_result = await conn.fetchrow("SELECT COUNT(*) as total FROM documents")
            total_count = total_count_result["total"] if total_count_result else 0
        
        return {
            "documents": documents,
            "total": total_count,
            "limit": limit,
            "offset": offset
        }
        
    except Exception as e:
        logger.error(f"Document listing failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.delete("/documents/{document_id}")
async def delete_document_endpoint(document_id: str):
    """Delete document endpoint."""
    try:
        # Validate UUID format
        try:
            uuid.UUID(document_id)
        except ValueError:
            raise HTTPException(status_code=400, detail="Invalid document ID format")
        
        success = await delete_document(document_id)
        
        if not success:
            raise HTTPException(status_code=404, detail="Document not found")
        
        return {
            "success": True,
            "message": f"Document {document_id} deleted successfully"
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Document deletion failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))

def _extract_text_from_document(doc: Dict[str, Any]) -> Optional[str]:
    """Best-effort text extraction from a document dict."""
    if not isinstance(doc, dict):
        return None

    # Try direct fields
    direct_candidates = [
        doc.get("content"),
        doc.get("text"),
        doc.get("plain_text"),
        doc.get("plaintext"),
        doc.get("full_text"),
        doc.get("raw_text"),
        doc.get("body"),
        doc.get("markdown"),
        doc.get("md"),
        doc.get("snippet"),
        doc.get("preview"),
        (doc.get("stats") or {}).get("content"),
        (doc.get("stats") or {}).get("text"),
        (doc.get("metadata") or {}).get("content"),
        (doc.get("metadata") or {}).get("text"),
        (doc.get("metadata") or {}).get("raw_text"),
        (doc.get("metadata") or {}).get("markdown"),
    ]
    for c in direct_candidates:
        if isinstance(c, str) and c.strip():
            return c

    # If content looks like JSON string, parse for common fields
    try:
        if isinstance(doc.get("content"), str) and doc["content"].strip().startswith("{"):
            data = json.loads(doc["content"])  # type: ignore
            for k in ("content", "text", "markdown", "md", "body"):
                v = data.get(k)
                if isinstance(v, str) and v.strip():
                    return v
    except Exception:
        pass

    # Try chunks
    chunks = doc.get("chunks")
    if isinstance(chunks, list) and chunks:
        parts: List[str] = []
        for ch in chunks[:50]:
            if not isinstance(ch, dict):
                continue
            t = ch.get("content") or ch.get("text") or ch.get("body")
            if isinstance(t, str) and t.strip():
                parts.append(t.strip())
            if len("\n\n".join(parts)) > 8000:
                break
        if parts:
            return "\n\n".join(parts)

    # Try pages array
    pages = doc.get("pages")
    if isinstance(pages, list) and pages:
        parts: List[str] = []
        for p in pages[:20]:
            if not isinstance(p, dict):
                continue
            t = p.get("content") or p.get("text")
            if isinstance(t, str) and t.strip():
                parts.append(t.strip())
            if len("\n\n".join(parts)) > 8000:
                break
        if parts:
            return "\n\n".join(parts)

    return None


@app.get("/documents/{document_id}")
async def get_document_endpoint(document_id: str):
    """Return a complete document with chunks for preview/enrichment."""
    try:
        input_data = DocumentInput(document_id=document_id)
        document = await get_document_tool(input_data)
        if not document:
            raise HTTPException(status_code=404, detail="Document not found")
        return document
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Get document failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/documents/{document_id}/{variant}")
async def get_document_content_endpoint(document_id: str, variant: str):
    """Return raw textual content for a document in various variants for preview."""
    allowed = {
        "content", "raw", "text", "plaintext", "plain", "body",
        "markdown", "md", "preview", "download", "file"
    }
    if variant not in allowed:
        raise HTTPException(status_code=404, detail="Unknown variant")

    try:
        input_data = DocumentInput(document_id=document_id)
        document = await get_document_tool(input_data)
        if not document:
            raise HTTPException(status_code=404, detail="Document not found")

        text = _extract_text_from_document(document)
        if not text or not text.strip():
            raise HTTPException(status_code=404, detail="No textual content available for preview")

        # Decide media type
        title = (document.get("title") or "").lower() if isinstance(document, dict) else ""
        source = (document.get("source") or "").lower() if isinstance(document, dict) else ""
        is_markdown = variant in {"markdown", "md"} or title.endswith(".md") or source.endswith(".md")
        media_type = "text/markdown; charset=utf-8" if is_markdown else "text/plain; charset=utf-8"
        return Response(content=text, media_type=media_type)
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Get document content failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/documents/{document_id}/summary")
async def generate_dbr_summary(
    document_id: str,
    request: Request
):
    """
    Generate a comprehensive summary of a DBR document using the RAG system.
    
    Args:
        document_id: UUID of the document to summarize
        request: Optional configuration for summary generation
    """
    try:
        from .summarizer import dbr_summarizer
        
        # Parse request body
        try:
            body = await request.json() if request else {}
        except:
            body = {}
            
        include_context = body.get("include_context", True)
        context_queries = body.get("context_queries")
        summary_type = body.get("summary_type", "comprehensive")
        force_regenerate = body.get("force_regenerate", False)
        
        # Validate summary type
        valid_types = {"comprehensive", "executive", "financial", "operational"}
        if summary_type not in valid_types:
            raise HTTPException(
                status_code=400,
                detail=f"Invalid summary_type. Must be one of: {', '.join(valid_types)}"
            )
        
        cache_action = "regenerating" if force_regenerate else "checking cache for"
        logger.info(f"Processing {summary_type} summary for document {document_id} ({cache_action})")
        
        # Generate summary (with caching logic)
        summary_result = await dbr_summarizer.summarize_dbr(
            document_id=document_id,
            include_context=include_context,
            context_queries=context_queries,
            summary_type=summary_type,
            force_regenerate=force_regenerate
        )
        
        return summary_result
        
    except ValueError as e:
        raise HTTPException(status_code=404, detail=str(e))
    except Exception as e:
        logger.error(f"DBR summary generation failed: {e}")
        raise HTTPException(status_code=500, detail=f"Summary generation failed: {str(e)}")


@app.post("/documents/{document_id}/summary_async")
async def start_dbr_summary_job(
    document_id: str,
    request: Request
):
    """
    Start an asynchronous summary generation job and return a job_id immediately.
    This avoids long-running HTTP requests that can time out on proxies.
    """
    try:
        from .summarizer import dbr_summarizer

        try:
            body = await request.json() if request else {}
        except Exception:
            body = {}

        include_context = body.get("include_context", True)
        context_queries = body.get("context_queries")
        summary_type = body.get("summary_type", "comprehensive")
        force_regenerate = body.get("force_regenerate", False)

        job_id = await create_summary_job(document_id, summary_type)

        async def _run_job():
            await update_summary_job_status(job_id, "running")
            try:
                result = await dbr_summarizer.summarize_dbr(
                    document_id=document_id,
                    include_context=include_context,
                    context_queries=context_queries,
                    summary_type=summary_type,
                    force_regenerate=force_regenerate,
                    job_id=job_id,
                )
                await set_summary_job_result(job_id, result)
            except Exception as e:
                logger.exception("Async summary job failed: %s", e)
                # If the job was cancelled, keep status as 'cancelled'
                try:
                    if await is_summary_job_cancelled(job_id):
                        await update_summary_job_status(job_id, "cancelled", error=str(e))
                    else:
                        await update_summary_job_status(job_id, "error", error=str(e))
                except Exception:
                    await update_summary_job_status(job_id, "error", error=str(e))

        # Fire-and-forget
        asyncio.create_task(_run_job())

        return JSONResponse(status_code=202, content={
            "job_id": job_id,
            "status": "queued",
            "document_id": document_id,
            "summary_type": summary_type,
        })
    except Exception as e:
        logger.error(f"Failed to start summary job: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to start job: {str(e)}")


@app.get("/jobs/{job_id}/status")
async def get_job_status(job_id: str):
    job = await get_summary_job(job_id)
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    # Do not include the full result here to keep polling payload small
    job.pop("result", None)
    # Normalize keys for client compatibility
    job["job_id"] = job.pop("id")
    return job


@app.get("/jobs/{job_id}/result")
async def get_job_result(job_id: str):
    job = await get_summary_job(job_id)
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    if job.get("status") != "done":
        return JSONResponse(status_code=202, content={
            "job_id": job_id,
            "status": job.get("status"),
            "error": job.get("error"),
        })
    return {
        "job_id": job.get("id", job_id),
        "status": job.get("status"),
        "result": job.get("result"),
    }


@app.post("/jobs/{job_id}/cancel")
async def cancel_job(job_id: str):
    ok = await cancel_summary_job(job_id)
    if not ok:
        raise HTTPException(status_code=404, detail="Job not found or cannot cancel")
    return {"job_id": job_id, "cancelled": True}


@app.get("/documents/{document_id}/summary/{summary_type}")
async def get_dbr_summary_typed(document_id: str, summary_type: str):
    """
    Generate a typed summary of a DBR document.
    Convenience endpoint for specific summary types.
    """
    from fastapi import Request
    import json
    
    # Create a mock request object with the summary type
    class MockRequest:
        async def json(self):
            return {"summary_type": summary_type, "include_context": True}
    
    return await generate_dbr_summary(document_id, MockRequest())


@app.get("/documents/{document_id}/summaries")
async def list_cached_summaries(document_id: str):
    """
    List all cached summaries for a document.
    """
    try:
        from .db_utils import list_document_summaries
        
        summaries = await list_document_summaries(document_id)
        return {
            "document_id": document_id,
            "cached_summaries": summaries,
            "total_count": len(summaries)
        }
        
    except Exception as e:
        logger.error(f"Failed to list cached summaries: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to list summaries: {str(e)}")


@app.delete("/documents/{document_id}/summaries")
async def clear_document_summaries(document_id: str, summary_type: str = None):
    """
    Clear cached summaries for a document.
    Query parameter summary_type can specify which type to clear, or clear all if omitted.
    """
    try:
        from .db_utils import delete_summary
        
        success = await delete_summary(document_id, summary_type)
        if success:
            message = f"Cleared {summary_type or 'all'} summaries for document {document_id}"
            return {"message": message, "success": True}
        else:
            raise HTTPException(status_code=500, detail="Failed to clear summaries")
            
    except Exception as e:
        logger.error(f"Failed to clear summaries: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to clear summaries: {str(e)}")


@app.get("/summaries/statistics")
async def get_summary_cache_statistics():
    """
    Get statistics about the summary cache.
    """
    try:
        from .db_utils import get_summary_statistics
        
        stats = await get_summary_statistics()
        return {
            "cache_statistics": stats,
            "cache_enabled": True
        }
        
    except Exception as e:
        logger.error(f"Failed to get summary statistics: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to get statistics: {str(e)}")


@app.post("/upload")
async def upload_document(file: UploadFile = File(...)):
    """Upload and ingest document endpoint."""
    if not INGESTION_AVAILABLE:
        raise HTTPException(
            status_code=503,
            detail="Document ingestion pipeline is not available"
        )
    
    try:
        # Validate file type
        allowed_extensions = {'.txt', '.md', '.pdf', '.doc', '.docx', '.xls', '.xlsx', '.csv', '.tsv'}
        file_extension = os.path.splitext(file.filename)[1].lower() if file.filename else ''
        
        if file_extension not in allowed_extensions:
            raise HTTPException(
                status_code=400, 
                detail=f"File type {file_extension} not supported. Allowed: {', '.join(allowed_extensions)}"
            )
        
        # Stream upload to disk with configurable size limit (default 200MB)
        try:
            max_mb = int(os.getenv("MAX_UPLOAD_MB", "200"))
        except Exception:
            max_mb = 200
        max_size = max_mb * 1024 * 1024
        
        # Create temporary directory for processing
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_file_path = Path(temp_dir) / file.filename
            
            # Write the uploaded file to disk in chunks to avoid high memory usage
            size_bytes = 0
            with open(temp_file_path, 'wb') as temp_file:
                chunk_size = 8 * 1024 * 1024  # 8MB
                while True:
                    chunk = await file.read(chunk_size)
                    if not chunk:
                        break
                    size_bytes += len(chunk)
                    if size_bytes > max_size:
                        raise HTTPException(
                            status_code=413,
                            detail=f"File too large. Maximum size is {max_mb}MB, got {size_bytes / 1024 / 1024:.1f}MB"
                        )
                    temp_file.write(chunk)

            # Convert to Markdown using converter (handles pdf/docx/xlsx/etc.)
            try:
                md_text, conv_meta = convert_to_markdown(str(temp_file_path))
            except Exception as e:
                logger.exception("Conversion to markdown failed: %s", e)
                raise HTTPException(status_code=500, detail="Failed to convert file to markdown")
            if not md_text or not md_text.strip():
                raise HTTPException(status_code=400, detail="No extractable text content found in the uploaded file")

            # Write markdown to a file that the ingestion pipeline will pick up
            md_name = f"{Path(file.filename).stem}.md"
            md_path = Path(temp_dir) / md_name
            with open(md_path, 'w', encoding='utf-8') as f_md:
                f_md.write(md_text)
            
            # Pre-ingestion diagnostics: list markdown files in staging
            found_md = sorted(glob.glob(str(Path(temp_dir) / "**/*.md"), recursive=True))

            # Create ingestion configuration (only supported fields)
            config = IngestionConfig(
                chunk_size=int(os.getenv("CHUNK_SIZE", "800")),
                chunk_overlap=int(os.getenv("CHUNK_OVERLAP", "150")),
                max_chunk_size=int(os.getenv("MAX_CHUNK_SIZE", "1500")),
                use_semantic_splitting=os.getenv("USE_SEMANTIC_SPLITTING", "1") not in ("0", "false", "False"),
                extract_entities=os.getenv("EXTRACT_ENTITIES", "1") not in ("0", "false", "False"),
                skip_graph_building=os.getenv("INGESTION_SKIP_GRAPH_BUILDING", "0") in ("1", "true", "True"),
            )
            
            # Initialize ingestion pipeline
            pipeline = DocumentIngestionPipeline(
                config=config,
                documents_folder=str(temp_dir),
                clean_before_ingest=False
            )
            
            # Process the document with optional global timeout
            try:
                global_timeout = float(os.getenv("INGEST_GLOBAL_TIMEOUT", "0"))
            except Exception:
                global_timeout = 0.0
            if global_timeout > 0:
                results = await asyncio.wait_for(pipeline.ingest_documents(), timeout=global_timeout)
            else:
                results = await pipeline.ingest_documents()

            # Check if results were successful (successful if document_id is not empty and no errors)
            successful_results = [r for r in results if r.document_id and len(r.errors) == 0]
            failed_results = [r for r in results if not r.document_id or len(r.errors) > 0]

            if not successful_results:
                # Return diagnostics to help identify conversion/ingestion issues
                diag = {
                    "message": "Document processing failed: No results returned",
                    "converted_markdown_chars": len(md_text or ""),
                    "found_markdown_files": found_md,
                    "ingestion_errors": [err for r in failed_results for err in (r.errors or [])],
                }
                logger.warning("Upload diagnostics: %s", diag)
                raise HTTPException(status_code=500, detail=diag)
            
            # Aggregate statistics from all successful results
            total_chunks = sum(r.chunks_created for r in successful_results)
            total_entities = sum(r.entities_extracted for r in successful_results)
            total_relationships = sum(r.relationships_created for r in successful_results)
            total_processing_time = sum(r.processing_time_ms for r in successful_results)
            
            logger.info(f"Successfully ingested document: {file.filename}")
            
            return {
                "success": True,
                "message": f"File {file.filename} uploaded and processed successfully",
                "filename": file.filename,
                "size": size_bytes,
                "converted_markdown_chars": len(md_text or ""),
                "found_markdown_files": found_md,
                "documents_processed": len(successful_results),
                "chunks_created": total_chunks,
                "embeddings_created": total_chunks,  # Each chunk gets an embedding
                "graph_entities": total_entities,
                "graph_relationships": total_relationships,
                "processing_time_ms": total_processing_time,
                "failed_documents": len(failed_results)
            }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"File upload and ingestion failed: {e}")
        raise HTTPException(status_code=500, detail=f"Upload failed: {str(e)}")


@app.post("/upload_async")
async def upload_document_async(file: UploadFile = File(...)):
    """Asynchronous upload endpoint.

    Saves the uploaded file to a staging directory and immediately returns 202,
    then converts + ingests in a background task to avoid long-running requests
    that can time out at proxies (e.g., Replit gateway).
    """
    if not INGESTION_AVAILABLE:
        raise HTTPException(
            status_code=503,
            detail="Document ingestion pipeline is not available"
        )

    try:
        allowed_extensions = {'.txt', '.md', '.pdf', '.doc', '.docx', '.xls', '.xlsx', '.csv', '.tsv'}
        original_name = file.filename or "upload.bin"
        file_extension = os.path.splitext(original_name)[1].lower()

        if file_extension not in allowed_extensions:
            raise HTTPException(
                status_code=400,
                detail=f"File type {file_extension} not supported. Allowed: {', '.join(allowed_extensions)}"
            )

        try:
            max_mb = int(os.getenv("MAX_UPLOAD_MB", "200"))
        except Exception:
            max_mb = 200
        max_size = max_mb * 1024 * 1024

        # Persist staging dir across background task lifecycle
        staging_dir = tempfile.mkdtemp(prefix="rag_upload_")
        temp_file_path = Path(staging_dir) / original_name

        # Stream write uploaded file
        size_bytes = 0
        with open(temp_file_path, 'wb') as temp_file:
            chunk_size = 8 * 1024 * 1024  # 8MB
            while True:
                chunk = await file.read(chunk_size)
                if not chunk:
                    break
                size_bytes += len(chunk)
                if size_bytes > max_size:
                    # Cleanup and abort
                    try:
                        shutil.rmtree(staging_dir, ignore_errors=True)
                    except Exception:
                        pass
                    raise HTTPException(
                        status_code=413,
                        detail=f"File too large. Maximum size is {max_mb}MB, got {size_bytes / 1024 / 1024:.1f}MB"
                    )
                temp_file.write(chunk)

        # Create an ingest job and return job_id immediately
        job_id = str(uuid.uuid4())
        INGEST_JOBS[job_id] = {
            "status": "queued",
            "error": None,
            "result": None,
            "filename": original_name,
            "size": size_bytes,
            "stage": "queued",
            "created_at": datetime.utcnow().isoformat(),
            "updated_at": datetime.utcnow().isoformat(),
        }

        async def _process_async_upload():
            """Background task to convert to markdown and ingest, then cleanup."""
            try:
                # Mark running
                try:
                    job = INGEST_JOBS.get(job_id)
                    if job is not None:
                        job["status"] = "running"
                        job["updated_at"] = datetime.utcnow().isoformat()
                except Exception:
                    pass

                # Stage: converting
                try:
                    job = INGEST_JOBS.get(job_id)
                    if job is not None:
                        job["stage"] = "converting"
                        job["updated_at"] = datetime.utcnow().isoformat()
                except Exception:
                    pass
                # Convert to Markdown
                try:
                    md_text, conv_meta = convert_to_markdown(str(temp_file_path))
                except Exception as e:
                    logger.exception("Conversion to markdown failed (async): %s", e)
                    try:
                        job = INGEST_JOBS.get(job_id)
                        if job is not None:
                            job["status"] = "error"
                            job["error"] = f"conversion_failed: {e}"
                            job["updated_at"] = datetime.utcnow().isoformat()
                    except Exception:
                        pass
                    return
                if not md_text or not md_text.strip():
                    logger.warning("No extractable text found in async upload: %s", original_name)
                    try:
                        job = INGEST_JOBS.get(job_id)
                        if job is not None:
                            job["status"] = "error"
                            job["error"] = "no_text_extracted"
                            job["updated_at"] = datetime.utcnow().isoformat()
                    except Exception:
                        pass
                    return

                md_name = f"{Path(original_name).stem}.md"
                md_path = Path(staging_dir) / md_name
                with open(md_path, 'w', encoding='utf-8') as f_md:
                    f_md.write(md_text)

                # Stage: converted
                try:
                    job = INGEST_JOBS.get(job_id)
                    if job is not None:
                        job["stage"] = "converted"
                        job["updated_at"] = datetime.utcnow().isoformat()
                except Exception:
                    pass

                # Pre-ingestion diagnostics
                found_md = sorted(glob.glob(str(Path(staging_dir) / "**/*.md"), recursive=True))

                # Ingestion config
                config = IngestionConfig(
                    chunk_size=int(os.getenv("CHUNK_SIZE", 800)),
                    chunk_overlap=int(os.getenv("CHUNK_OVERLAP", 150)),
                    max_chunk_size=int(os.getenv("MAX_CHUNK_SIZE", 1500)),
                    embedding_model=os.getenv("EMBEDDING_MODEL", "text-embedding-004"),
                    vector_dimension=int(os.getenv("VECTOR_DIMENSION", 768)),
                    llm_choice=os.getenv("INGESTION_LLM_CHOICE", os.getenv("LLM_CHOICE", "gemini-1.5-flash")),
                    enable_graph=True,
                    clean_before_ingest=False,
                )

                pipeline = DocumentIngestionPipeline(
                    config=config,
                    documents_folder=staging_dir,
                    clean_before_ingest=False,
                )

                results = await pipeline.ingest_documents()
                successful_results = [r for r in results if r.document_id and len(r.errors) == 0]
                failed_results = [r for r in results if not r.document_id or len(r.errors) > 0]

                if successful_results:
                    logger.info(
                        "Async upload ingested %d document(s) (chunks=%d) for %s",
                        len(successful_results),
                        sum(r.chunks_created for r in successful_results),
                        original_name,
                    )
                else:
                    logger.warning("Async upload produced no successful results for %s; errors=%s; found_md=%s",
                                   original_name,
                                   [err for r in failed_results for err in (r.errors or [])],
                                   found_md)
            except Exception as e:
                logger.exception("Async upload processing failed: %s", e)
            finally:
                # Cleanup staging dir
                try:
                    shutil.rmtree(staging_dir, ignore_errors=True)
                except Exception:
                    pass

        # Kick off background processing and return immediately
        asyncio.create_task(_process_async_upload())

        return JSONResponse(status_code=202, content={
            "success": True,
            "status": "queued",
            "filename": original_name,
            "size": size_bytes,
            "job_id": job_id,
        })

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"File async upload scheduling failed: {e}")
        raise HTTPException(status_code=500, detail=f"Upload scheduling failed: {str(e)}")


@app.post("/uploads/initiate")
async def initiate_multipart_upload(request: Request):
    """Initiate a multipart upload session for large files.

    Body JSON:
      - filename: original filename
      - total_parts: expected number of chunks (int)
    """
    try:
        body = await request.json()
        filename = (body.get("filename") or "upload.bin").strip()
        total_parts = int(body.get("total_parts"))
        if not filename or total_parts <= 0:
            raise HTTPException(status_code=400, detail="filename and total_parts are required")

        # Create staging dir for this session
        staging_dir = tempfile.mkdtemp(prefix="rag_mpu_")
        upload_id = str(uuid.uuid4())
        UPLOAD_SESSIONS[upload_id] = {
            "filename": filename,
            "total_parts": total_parts,
            "staging_dir": staging_dir,
            "created_at": datetime.utcnow().isoformat(),
        }
        return {"upload_id": upload_id}
    except HTTPException:
        raise
    except Exception as e:
        logger.exception("Failed to initiate multipart upload: %s", e)
        raise HTTPException(status_code=500, detail="Failed to initiate upload")


@app.post("/uploads/{upload_id}/part")
async def upload_multipart_part(upload_id: str, index: int = Form(...), chunk: UploadFile = File(...)):
    """Upload a single part for a multipart upload session.

    Accepts multipart/form-data with fields:
      - index: 0-based part index
      - chunk: file part payload
    """
    session = UPLOAD_SESSIONS.get(upload_id)
    if not session:
        raise HTTPException(status_code=404, detail="Upload session not found")
    try:
        staging_dir = session["staging_dir"]
        total_parts = int(session.get("total_parts") or 0)
        if index < 0 or (total_parts and index >= total_parts):
            raise HTTPException(status_code=400, detail="Invalid part index")

        part_path = Path(staging_dir) / f"part_{index:06d}"
        # Write part to disk
        size_bytes = 0
        with open(part_path, "wb") as f:
            while True:
                data = await chunk.read(8 * 1024 * 1024)
                if not data:
                    break
                size_bytes += len(data)
                f.write(data)
        return {"ok": True, "index": index, "size": size_bytes}
    except HTTPException:
        raise
    except Exception as e:
        logger.exception("Failed to write multipart part: %s", e)
        raise HTTPException(status_code=500, detail="Failed to write part")


@app.post("/uploads/{upload_id}/complete")
async def complete_multipart_upload(upload_id: str):
    """Complete multipart upload: stitch parts and start async ingestion."""
    session = UPLOAD_SESSIONS.get(upload_id)
    if not session:
        raise HTTPException(status_code=404, detail="Upload session not found")
    staging_dir = session["staging_dir"]
    filename = session["filename"]
    total_parts = int(session.get("total_parts") or 0)

    try:
        staging = Path(staging_dir)
        # Validate parts
        parts = sorted(staging.glob("part_*"))
        if total_parts and len(parts) != total_parts:
            raise HTTPException(status_code=400, detail=f"Expected {total_parts} parts, got {len(parts)}")

        assembled_path = staging / filename
        # Stitch parts
        with open(assembled_path, "wb") as out_f:
            for p in parts:
                with open(p, "rb") as in_f:
                    shutil.copyfileobj(in_f, out_f, length=8 * 1024 * 1024)

        # Background processing similar to /upload_async
        # Create an ingest job for this session
        job_id = str(uuid.uuid4())
        INGEST_JOBS[job_id] = {
            "status": "queued",
            "error": None,
            "result": None,
            "filename": filename,
            "stage": "queued",
            "created_at": datetime.utcnow().isoformat(),
            "updated_at": datetime.utcnow().isoformat(),
        }

        async def _process_assembled():
            try:
                # Mark running
                try:
                    job = INGEST_JOBS.get(job_id)
                    if job is not None:
                        job["status"] = "running"
                        job["updated_at"] = datetime.utcnow().isoformat()
                except Exception:
                    pass
                # Stage: converting
                try:
                    job = INGEST_JOBS.get(job_id)
                    if job is not None:
                        job["stage"] = "converting"
                        job["updated_at"] = datetime.utcnow().isoformat()
                except Exception:
                    pass
                # Convert to markdown
                md_text, conv_meta = convert_to_markdown(str(assembled_path))
                if not md_text or not md_text.strip():
                    logger.warning("No extractable text in assembled file %s", assembled_path)
                    try:
                        job = INGEST_JOBS.get(job_id)
                        if job is not None:
                            job["status"] = "error"
                            job["stage"] = "error"
                            job["error"] = "no_text_extracted"
                            job["updated_at"] = datetime.utcnow().isoformat()
                    except Exception:
                        pass
                    return
                md_path = staging / f"{Path(filename).stem}.md"
                with open(md_path, "w", encoding="utf-8") as f_md:
                    f_md.write(md_text)

                config = IngestionConfig(
                    chunk_size=int(os.getenv("CHUNK_SIZE", 800)),
                    chunk_overlap=int(os.getenv("CHUNK_OVERLAP", 150)),
                    max_chunk_size=int(os.getenv("MAX_CHUNK_SIZE", 1500)),
                    embedding_model=os.getenv("EMBEDDING_MODEL", "text-embedding-004"),
                    vector_dimension=int(os.getenv("VECTOR_DIMENSION", 768)),
                    llm_choice=os.getenv("INGESTION_LLM_CHOICE", os.getenv("LLM_CHOICE", "gemini-1.5-flash")),
                    enable_graph=True,
                    clean_before_ingest=False,
                )
                pipeline = DocumentIngestionPipeline(
                    config=config,
                    documents_folder=staging_dir,
                    clean_before_ingest=False,
                )
                # Stage: ingesting
                try:
                    job = INGEST_JOBS.get(job_id)
                    if job is not None:
                        job["stage"] = "ingesting"
                        job["updated_at"] = datetime.utcnow().isoformat()
                except Exception:
                    pass

                results = await pipeline.ingest_documents()
                ok = [r for r in results if getattr(r, "document_id", "") and len(getattr(r, "errors", []) or []) == 0]
                try:
                    job = INGEST_JOBS.get(job_id)
                    if job is not None:
                        job["status"] = "done"
                        job["stage"] = "done"
                        job["result"] = {"document_ids": [r.document_id for r in ok if r.document_id]}
                        job["updated_at"] = datetime.utcnow().isoformat()
                except Exception:
                    pass
            except Exception as e:
                logger.exception("Multipart completion failed: %s", e)
                try:
                    job = INGEST_JOBS.get(job_id)
                    if job is not None:
                        job["status"] = "error"
                        job["stage"] = "error"
                        job["error"] = str(e)
                        job["updated_at"] = datetime.utcnow().isoformat()
                except Exception:
                    pass
            finally:
                # Cleanup session and staging
                try:
                    shutil.rmtree(staging_dir, ignore_errors=True)
                finally:
                    UPLOAD_SESSIONS.pop(upload_id, None)

        asyncio.create_task(_process_assembled())
        return JSONResponse(status_code=202, content={"success": True, "status": "queued", "job_id": job_id})
    except HTTPException:
        raise
    except Exception as e:
        logger.exception("Failed to complete multipart upload: %s", e)
        raise HTTPException(status_code=500, detail="Failed to complete upload")


@app.get("/ingest/jobs/{job_id}/status")
async def get_ingest_job_status(job_id: str):
    """Return current status of an ingestion job created by upload_async or multipart complete."""
    job = INGEST_JOBS.get(job_id)
    if not job:
        raise HTTPException(status_code=404, detail="job_not_found")
    # Return a copy to avoid accidental mutation
    return {
        "job_id": job_id,
        "status": job.get("status"),
        "stage": job.get("stage"),
        "error": job.get("error"),
        "result": job.get("result"),
        "filename": job.get("filename"),
        "size": job.get("size"),
        "created_at": job.get("created_at"),
        "updated_at": job.get("updated_at"),
    }


@app.get("/ingest/jobs/{job_id}/result")
async def get_ingest_job_result(job_id: str):
    """Return job result when done. Returns 202 while still processing."""
    job = INGEST_JOBS.get(job_id)
    if not job:
        raise HTTPException(status_code=404, detail="job_not_found")
    status = (job.get("status") or "").lower()
    if status == "done":
        return {
            "job_id": job_id,
            "status": status,
            "result": job.get("result") or {},
        }
    if status == "error":
        return JSONResponse(status_code=500, content={
            "job_id": job_id,
            "status": status,
            "error": job.get("error") or "unknown_error",
        })
    # queued or running
    return JSONResponse(status_code=202, content={
        "job_id": job_id,
        "status": status or "queued",
    })


@app.get("/ingest/jobs")
async def list_ingest_jobs(status: Optional[str] = None):
    """List ingestion jobs (ephemeral, in-memory). Optional filter by status.

    Query params:
      - status: queued | running | done | error
    """
    try:
        items = []
        for jid, job in INGEST_JOBS.items():
            if status and (job.get("status") or "").lower() != status.lower():
                continue
            items.append({
                "job_id": jid,
                "status": job.get("status"),
                "stage": job.get("stage"),
                "error": job.get("error"),
                "result": job.get("result"),
                "filename": job.get("filename"),
                "size": job.get("size"),
                "created_at": job.get("created_at"),
                "updated_at": job.get("updated_at"),
            })
        # Sort newest first by created_at when available
        try:
            items.sort(key=lambda x: x.get("created_at") or "", reverse=True)
        except Exception:
            pass
        return {"jobs": items, "count": len(items)}
    except Exception as e:
        logger.error(f"Failed to list ingest jobs: {e}")
        raise HTTPException(status_code=500, detail="Failed to list ingest jobs")


@app.post("/convert")
async def convert_only(file: UploadFile = File(...)):
    """Server-side conversion to Markdown without ingestion.

    Returns JSON with diagnostics: original ext, markdown length, and preview.
    Useful on Replit to verify converters (python-docx, openpyxl, pdfminer/ocr) work.
    """
    if not file.filename:
        raise HTTPException(status_code=400, detail="Filename is required")

    try:
        # Validate extension (same as /upload)
        allowed_extensions = {'.txt', '.md', '.pdf', '.doc', '.docx', '.xls', '.xlsx', '.csv', '.tsv'}
        ext = os.path.splitext(file.filename)[1].lower()
        if ext not in allowed_extensions:
            raise HTTPException(status_code=400, detail=f"Unsupported extension {ext}")

        # Size limit
        try:
            max_mb = int(os.getenv("MAX_UPLOAD_MB", "200"))
        except Exception:
            max_mb = 200
        max_size = max_mb * 1024 * 1024

        with tempfile.TemporaryDirectory() as temp_dir:
            temp_file_path = Path(temp_dir) / file.filename
            size_bytes = 0
            with open(temp_file_path, 'wb') as temp_file:
                chunk_size = 8 * 1024 * 1024
                while True:
                    chunk = await file.read(chunk_size)
                    if not chunk:
                        break
                    size_bytes += len(chunk)
                    if size_bytes > max_size:
                        raise HTTPException(status_code=413, detail=f"File too large (> {max_mb}MB)")
                    temp_file.write(chunk)

            # Convert to markdown
            md_text, meta = convert_to_markdown(str(temp_file_path))
            md_text = md_text or ""
            preview = md_text[:1000]
            ocr_enabled = os.getenv("OCR_PDF", "0").lower() in {"1", "true", "yes", "on"}
            return {
                "success": True,
                "filename": file.filename,
                "original_ext": ext,
                "size": size_bytes,
                "converted_markdown_chars": len(md_text),
                "preview": preview,
                "ocr_enabled": ocr_enabled,
                "meta": meta,
            }
    except HTTPException:
        raise
    except Exception as e:
        logger.exception("Conversion failed: %s", e)
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/ingest/folder")
async def ingest_folder(request: Request):
    """Bulk-ingest a local folder. Converts supported files to Markdown first, then ingests.

    Body JSON:
      - path: absolute path to a folder containing source files
      - collection_id (optional): add ingested documents to this collection
    """
    if not INGESTION_AVAILABLE:
        raise HTTPException(status_code=503, detail="Document ingestion pipeline is not available")

    try:
        body = await request.json()
        folder_path = body.get("path")
        collection_id = body.get("collection_id")
        if not folder_path or not isinstance(folder_path, str):
            raise HTTPException(status_code=400, detail="'path' is required")
        src = Path(folder_path)
        if not src.exists() or not src.is_dir():
            raise HTTPException(status_code=400, detail="Provided path does not exist or is not a directory")

        allowed_exts = {".txt", ".md", ".markdown", ".pdf", ".doc", ".docx", ".xls", ".xlsx", ".csv", ".tsv", ".html", ".htm", ".pptx"}

        with tempfile.TemporaryDirectory() as staging_dir:
            staging = Path(staging_dir)
            converted_count = 0
            errors: List[str] = []

            # Walk and convert
            for path in src.rglob("*"):
                if not path.is_file():
                    continue
                if path.suffix.lower() not in allowed_exts:
                    continue
                try:
                    text, meta = convert_to_markdown(str(path))
                    if not text or not text.strip():
                        errors.append(f"No text extracted: {path.name}")
                        continue
                    # write .md preserving relative structure
                    rel = path.relative_to(src)
                    out_path = staging / rel.with_suffix(".md")
                    out_path.parent.mkdir(parents=True, exist_ok=True)
                    with open(out_path, 'w', encoding='utf-8') as f:
                        f.write(text)
                    converted_count += 1
                except Exception as e:
                    logger.warning("Failed to convert %s: %s", path, e)
                    errors.append(f"{path.name}: {e}")

            if converted_count == 0:
                return {"success": False, "message": "No convertible files found", "errors": errors}

            # Build ingestion config
            config = IngestionConfig(
                chunk_size=int(os.getenv("CHUNK_SIZE", 800)),
                chunk_overlap=int(os.getenv("CHUNK_OVERLAP", 150)),
                max_chunk_size=int(os.getenv("MAX_CHUNK_SIZE", 1500)),
                embedding_model=os.getenv("EMBEDDING_MODEL", "text-embedding-004"),
                vector_dimension=int(os.getenv("VECTOR_DIMENSION", 768)),
                llm_choice=os.getenv("INGESTION_LLM_CHOICE", os.getenv("LLM_CHOICE", "gemini-1.5-flash")),
                enable_graph=True,
                clean_before_ingest=False
            )

            pipeline = DocumentIngestionPipeline(
                config=config,
                documents_folder=str(staging),
                clean_before_ingest=False
            )

            results = await pipeline.ingest_documents()
            successful = [r for r in results if r.document_id and len(r.errors) == 0]

            # Optionally add to collection
            added_to_collection = 0
            if collection_id and successful:
                try:
                    ids = [r.document_id for r in successful]
                    _ = await add_documents_to_collection_db(collection_id, ids, added_by=request.headers.get("x-user-id"))
                    added_to_collection = len(ids)
                except Exception as e:
                    logger.warning("Failed to add docs to collection %s: %s", collection_id, e)

            return {
                "success": True,
                "converted_files": converted_count,
                "documents_processed": len(successful),
                "errors": errors,
                "added_to_collection": added_to_collection,
            }
    except HTTPException:
        raise
    except Exception as e:
        logger.exception("Bulk ingest failed: %s", e)
        raise HTTPException(status_code=500, detail=str(e))


def _extract_section4(text: str) -> Optional[str]:
    """Heuristic extraction of Section 4 from a document's text."""
    if not text:
        return None
    # Normalize line endings
    t = text
    # Try to find a 'Section 4' header and capture until the next 'Section 5'
    patterns = [
        r"(?is)\bsection\s*4\b.*?(?=\n\s*section\s*5\b|\n\s*5\.|\n\s*section\s*V|\Z)",
        r"(?is)\n\s*4\.?\s+[A-Z].*?(?=\n\s*5\.|\n\s*section\s*5\b|\Z)",
    ]
    for pat in patterns:
        m = re.search(pat, t)
        if m:
            return m.group(0).strip()
    return None


@app.post("/reports/generate")
async def reports_generate(request: Request):
    """Generate a minimal report focusing on Section 4 (post-maps) for documents in a collection.

    Body JSON:
      - collection_id: required
      - limit (optional): limit number of docs processed
    """
    try:
        body = await request.json()
        collection_id = body.get("collection_id")
        limit = int(body.get("limit", 50))
        if not collection_id:
            raise HTTPException(status_code=400, detail="collection_id is required")

        # List documents in the collection
        docs, total = await list_collection_documents_db(collection_id, limit=limit, offset=0)
        from .tools import DocumentInput
        extracted = []
        for d in docs:
            doc_id = d.get("id") or d.get("document_id") or d.get("doc_id")
            title = d.get("title") or d.get("name") or "Untitled"
            if not doc_id:
                continue
            try:
                di = DocumentInput(document_id=str(doc_id))
                doc = await get_document_tool(di)
                text = _extract_text_from_document(doc) or ""
                section4 = _extract_section4(text) or ""
                if section4:
                    extracted.append({
                        "document_id": str(doc_id),
                        "title": title,
                        "section4": section4[:10000],  # cap
                        "citation": {
                            "document_title": title,
                            "match_index": max(text.lower().find(section4[:50].lower()), 0)
                        }
                    })
            except Exception as e:
                logger.warning("Failed to extract Section 4 for %s: %s", title, e)

        return {
            "collection_id": collection_id,
            "total_documents": total,
            "processed": len(docs),
            "extracted_count": len(extracted),
            "results": extracted,
        }
    except HTTPException:
        raise
    except Exception as e:
        logger.exception("Report generation failed: %s", e)
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/sessions/{session_id}")
async def get_session_info(session_id: str):
    """Get session information."""
    try:
        session = await get_session(session_id)
        if not session:
            raise HTTPException(status_code=404, detail="Session not found")
        
        return session
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Session retrieval failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# 3rd Party Integration Endpoints
_integration_instances = {}  # Store active integration instances

@app.get("/integrations")
async def list_available_integrations():
    """List available 3rd party integrations."""
    if not INTEGRATIONS_AVAILABLE:
        return {"integrations": []}
    
    integrations = [
        {
            "name": "google_drive",
            "display_name": "Google Drive",
            "description": "Import documents from Google Drive",
            "auth_required": True,
            "supported_formats": [".txt", ".md", ".pdf", ".docx", ".doc"]
        },
        {
            "name": "dropbox", 
            "display_name": "Dropbox",
            "description": "Import documents from Dropbox",
            "auth_required": True,
            "supported_formats": [".txt", ".md", ".pdf", ".docx", ".doc"]
        },
        {
            "name": "onedrive",
            "display_name": "OneDrive",
            "description": "Import documents from OneDrive", 
            "auth_required": True,
            "supported_formats": [".txt", ".md", ".pdf", ".docx", ".doc"]
        }
    ]
    
    return {"integrations": integrations}


@app.get("/integrations/{service_name}/auth-config")
async def get_integration_auth_config(service_name: str):
    """Get OAuth configuration for client-side authentication."""
    if not INTEGRATIONS_AVAILABLE:
        raise HTTPException(status_code=503, detail="Integrations not available")
    
    # Return public OAuth configuration (no secrets)
    configs = {
        "google_drive": {
            "client_id": os.getenv("GOOGLE_DRIVE_CLIENT_ID", "your_google_client_id"),  # Public client ID
            "auth_url": "https://accounts.google.com/o/oauth2/v2/auth",
            "scope": "https://www.googleapis.com/auth/drive.readonly https://www.googleapis.com/auth/drive.metadata.readonly",
            "redirect_uri": "http://localhost:3000/auth/callback/google_drive"
        },
        "dropbox": {
            "client_id": os.getenv("DROPBOX_CLIENT_ID", "your_dropbox_client_id"),  # Public client ID
            "auth_url": "https://www.dropbox.com/oauth2/authorize",
            "scope": "files.content.read files.metadata.read",
            "redirect_uri": "http://localhost:3000/auth/callback/dropbox"
        },
        "onedrive": {
            "client_id": os.getenv("ONEDRIVE_CLIENT_ID", "your_onedrive_client_id"),  # Public client ID
            "auth_url": "https://login.microsoftonline.com/common/oauth2/v2.0/authorize",
            "scope": "https://graph.microsoft.com/Files.Read https://graph.microsoft.com/Files.Read.All offline_access",
            "redirect_uri": "http://localhost:3000/auth/callback/onedrive"
        }
    }
    
    if service_name not in configs:
        raise HTTPException(status_code=400, detail=f"Unknown service: {service_name}")
    
    return configs[service_name]


@app.post("/integrations/{service_name}/store-token")
async def store_integration_token(service_name: str, request: Request):
    """Store user's OAuth token for integration."""
    if not INTEGRATIONS_AVAILABLE:
        raise HTTPException(status_code=503, detail="Integrations not available")
    
    try:
        body = await request.json()
        access_token = body.get('access_token')
        user_id = body.get('user_id', 'anonymous')  # Optional user identification
        
        if not access_token:
            raise HTTPException(status_code=400, detail="Missing access_token")
        
        # Create integration instance with user's token
        # No client secret needed for token-only operations
        if service_name == "google_drive":
            integration = GoogleDriveIntegration(IntegrationConfig(
                client_id="", client_secret="", redirect_uri="", 
                scopes=[], service_name=service_name
            ))
        elif service_name == "dropbox":
            integration = DropboxIntegration(IntegrationConfig(
                client_id="", client_secret="", redirect_uri="",
                scopes=[], service_name=service_name
            ))
        elif service_name == "onedrive":
            integration = OneDriveIntegration(IntegrationConfig(
                client_id="", client_secret="", redirect_uri="",
                scopes=[], service_name=service_name
            ))
        else:
            raise HTTPException(status_code=400, detail=f"Unknown service: {service_name}")
        
        # Set the user's token
        integration.set_tokens(access_token)
        
        # Store integration session
        session_id = str(uuid.uuid4())
        _integration_instances[session_id] = integration
        
        return {
            "success": True,
            "session_id": session_id,
            "service_name": service_name,
            "message": "Token stored successfully"
        }
        
    except Exception as e:
        logger.error(f"Token storage failed for {service_name}: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/integrations/{service_name}/documents")
async def list_integration_documents(
    service_name: str,
    session_id: str,
    folder_id: Optional[str] = None
):
    """List documents from a 3rd party integration."""
    if not INTEGRATIONS_AVAILABLE:
        raise HTTPException(status_code=503, detail="Integrations not available")
    
    integration = _integration_instances.get(session_id)
    if not integration:
        raise HTTPException(status_code=401, detail="Not authenticated or session expired")
    
    try:
        documents = await integration.list_documents(folder_id)
        return {"documents": documents}
        
    except Exception as e:
        logger.error(f"Failed to list {service_name} documents: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/integrations/{service_name}/import")
async def import_documents_from_integration(
    service_name: str,
    request: Request
):
    """Import selected documents from a 3rd party integration."""
    if not INTEGRATIONS_AVAILABLE:
        raise HTTPException(status_code=503, detail="Integrations not available")
    
    if not INGESTION_AVAILABLE:
        raise HTTPException(status_code=503, detail="Document ingestion not available")
    
    try:
        body = await request.json()
        session_id = body.get('session_id')
        document_ids = body.get('document_ids', [])
        
        if not session_id or not document_ids:
            raise HTTPException(status_code=400, detail="Missing session_id or document_ids")
        
        integration = _integration_instances.get(session_id)
        if not integration:
            raise HTTPException(status_code=401, detail="Not authenticated or session expired")
        
        imported_documents = []
        failed_documents = []
        total_documents = len(document_ids)
        processed_count = 0
        
        logger.info(f"Starting sequential import of {total_documents} documents from {service_name}")
        
        # Import each document sequentially (one at a time for data integrity)
        async for document in integration.bulk_import_documents(document_ids):
            processed_count += 1
            logger.info(f"Processing document {processed_count}/{total_documents}: {document.name}")
            
            try:
                # Save to temporary file
                temp_file_path = await integration.save_temp_file(document)
                
                # Create ingestion configuration
                config = IngestionConfig(
                    chunk_size=int(os.getenv("CHUNK_SIZE", 800)),
                    chunk_overlap=int(os.getenv("CHUNK_OVERLAP", 150)),
                    max_chunk_size=int(os.getenv("MAX_CHUNK_SIZE", 1500)),
                    embedding_model=os.getenv("EMBEDDING_MODEL", "text-embedding-004"),
                    vector_dimension=int(os.getenv("VECTOR_DIMENSION", 768)),
                    llm_choice=os.getenv("INGESTION_LLM_CHOICE", os.getenv("LLM_CHOICE", "gemini-1.5-flash")),
                    enable_graph=True,
                    clean_before_ingest=False
                )
                
                # Process with ingestion pipeline (sequential processing ensures data integrity)
                temp_dir = os.path.dirname(temp_file_path)
                pipeline = DocumentIngestionPipeline(
                    config=config,
                    documents_folder=temp_dir,
                    clean_before_ingest=False
                )
                
                logger.info(f"Starting ingestion for: {document.name}")
                results = await pipeline.ingest_documents()
                
                # Clean up temp file
                try:
                    os.unlink(temp_file_path)
                except:
                    pass  # Ignore cleanup errors
                
                # Process results
                successful_results = [r for r in results if r.document_id and len(r.errors) == 0]
                
                if successful_results:
                    result = successful_results[0]  # Should only be one document
                    imported_documents.append({
                        "id": document.id,
                        "name": document.name,
                        "service": service_name,
                        "chunks_created": result.chunks_created,
                        "entities_extracted": result.entities_extracted,
                        "processing_time_ms": result.processing_time_ms,
                        "processed_at": processed_count,
                        "total_documents": total_documents
                    })
                    logger.info(f"Successfully imported: {document.name} ({result.chunks_created} chunks, {result.entities_extracted} entities)")
                else:
                    failed_documents.append({
                        "id": document.id,
                        "name": document.name,
                        "error": "Processing failed - no successful results",
                        "processed_at": processed_count,
                        "total_documents": total_documents
                    })
                    logger.warning(f"Failed to import: {document.name} - no successful results")
                    
            except Exception as e:
                logger.error(f"Failed to import document {document.id}: {e}")
                failed_documents.append({
                    "id": document.id if hasattr(document, 'id') else 'unknown',
                    "name": document.name if hasattr(document, 'name') else 'unknown',
                    "error": str(e),
                    "processed_at": processed_count,
                    "total_documents": total_documents
                })
        
        return {
            "success": True,
            "imported_count": len(imported_documents),
            "failed_count": len(failed_documents),
            "imported_documents": imported_documents,
            "failed_documents": failed_documents
        }
    except Exception as e:
        logger.error(f"Bulk import from {service_name} failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))

# Analytics endpoints
@app.get("/api/analytics/real-time", response_model=RealTimeMetrics)
async def api_get_real_time_metrics():
    """Get real-time metrics."""
    metrics = await analytics_tracker.get_real_time_metrics()
    if metrics is None:
        metrics = RealTimeMetrics(
            active_sessions=0,
            messages_last_hour=0,
            new_users_last_hour=0,
            total_documents=0,
            documents_today=0,
            public_templates=0,
            total_collections=0,
        )
    return metrics


@app.get("/api/analytics/chat-metrics", response_model=ChatMetrics)
async def api_get_chat_metrics(days: int = 7):
    """Get aggregated chat metrics over the last N days."""
    metrics = await analytics_tracker.get_chat_activity_metrics(days=days)
    if metrics is None:
        metrics = ChatMetrics(
            total_messages=0,
            total_sessions=0,
            unique_users=0,
            avg_messages_per_session=0.0,
            total_tool_calls=0,
            avg_response_time_ms=0.0,
        )
    return metrics


@app.get("/api/analytics/document-usage", response_model=DocumentUsageStats)
async def api_get_document_usage():
    """Get document usage statistics."""
    stats = await analytics_tracker.get_document_usage_stats()
    if stats is None:
        stats = DocumentUsageStats(
            total_documents=0,
            documents_uploaded_today=0,
            most_referenced_document_id=None,
            most_referenced_document_title=None,
            avg_document_size=0,
        )
    return stats


@app.get("/api/analytics/trending-searches")
async def api_get_trending_searches(days: int = 7, limit: int = 10):
    """Get trending searches within the last N days."""
    searches = await analytics_tracker.get_trending_searches(days=days, limit=limit)
    return {"trending_searches": searches}


@app.get("/api/analytics/user-engagement", response_model=UserEngagementMetrics)
async def api_get_user_engagement(user_id: Optional[str] = None, days: int = 30):
    """Get user engagement metrics optionally filtered by user_id."""
    data = await analytics_tracker.get_user_engagement_metrics(user_id=user_id, days=days)
    # Provide safe defaults if empty
    return UserEngagementMetrics(
        total_sessions=int(data.get("total_sessions", 0) or 0),
        avg_messages_per_session=float(data.get("avg_messages_per_session", 0) or 0),
        total_messages=int(data.get("total_messages", 0) or 0),
        total_tool_calls=int(data.get("total_tool_calls", 0) or 0),
        total_searches=int(data.get("total_searches", 0) or 0),
        avg_response_time=float(data.get("avg_response_time", 0) or 0),
        high_satisfaction_count=int(data.get("high_satisfaction_count", 0) or 0),
        low_satisfaction_count=int(data.get("low_satisfaction_count", 0) or 0),
        avg_satisfaction_rating=float(data.get("avg_satisfaction_rating", 0) or 0),
    )


@app.get("/api/analytics/dashboard", response_model=AnalyticsDashboardResponse)
async def api_get_dashboard(days: int = 7, user_id: Optional[str] = None):
    """Get combined analytics dashboard data."""
    real_time = await analytics_tracker.get_real_time_metrics()
    if real_time is None:
        real_time = RealTimeMetrics(
            active_sessions=0,
            messages_last_hour=0,
            new_users_last_hour=0,
            total_documents=0,
            documents_today=0,
            public_templates=0,
            total_collections=0,
        )

    chat = await analytics_tracker.get_chat_activity_metrics(days=days)
    if chat is None:
        chat = ChatMetrics(
            total_messages=0,
            total_sessions=0,
            unique_users=0,
            avg_messages_per_session=0.0,
            total_tool_calls=0,
            avg_response_time_ms=0.0,
        )

    doc_stats = await analytics_tracker.get_document_usage_stats()
    if doc_stats is None:
        doc_stats = DocumentUsageStats(
            total_documents=0,
            documents_uploaded_today=0,
            most_referenced_document_id=None,
            most_referenced_document_title=None,
            avg_document_size=0,
        )

    trending = await analytics_tracker.get_trending_searches(days=days, limit=10)
    engagement_dict = await analytics_tracker.get_user_engagement_metrics(user_id=user_id, days=days)
    engagement = UserEngagementMetrics(
        total_sessions=int(engagement_dict.get("total_sessions", 0) or 0),
        avg_messages_per_session=float(engagement_dict.get("avg_messages_per_session", 0) or 0),
        total_messages=int(engagement_dict.get("total_messages", 0) or 0),
        total_tool_calls=int(engagement_dict.get("total_tool_calls", 0) or 0),
        total_searches=int(engagement_dict.get("total_searches", 0) or 0),
        avg_response_time=float(engagement_dict.get("avg_response_time", 0) or 0),
        high_satisfaction_count=int(engagement_dict.get("high_satisfaction_count", 0) or 0),
        low_satisfaction_count=int(engagement_dict.get("low_satisfaction_count", 0) or 0),
        avg_satisfaction_rating=float(engagement_dict.get("avg_satisfaction_rating", 0) or 0),
    )

    return AnalyticsDashboardResponse(
        real_time_metrics=real_time,
        chat_metrics=chat,
        document_stats=doc_stats,
        trending_searches=trending,
        user_engagement=engagement,
    )


# Question Generation Endpoints
@app.get("/api/questions/generate")
async def generate_questions(collection_id: Optional[str] = None, limit: int = 6):
    """Generate relevant questions based on document content."""
    logger.info(f"Question generation requested - collection_id: {collection_id}, limit: {limit}")
    
    if not question_generator:
        logger.warning("Question generator not initialized")
        return {"questions": [], "source": "fallback_no_generator", "error": "Question generator not initialized"}
    
    try:
        if collection_id:
            questions = await question_generator.generate_questions_for_collection(
                collection_id=collection_id, 
                limit=limit
            )
        else:
            questions = await question_generator.generate_questions_for_all_documents(
                limit=limit
            )
        
        logger.info(f"Generated {len(questions)} questions")
        return {"questions": questions, "source": "ai_generated"}
    
    except Exception as e:
        logger.error(f"Error generating questions: {e}")
        return {"questions": [], "source": "fallback_error", "error": str(e)}

@app.post("/api/questions/clear-cache")
async def clear_question_cache():
    """Clear the question generation cache."""
    if question_generator:
        question_generator.clear_cache()
        return {"message": "Question cache cleared"}
    return {"message": "Question generator not available"}

@app.get("/api/questions/debug")
async def debug_questions():
    """Debug endpoint to check document count and question generation."""
    try:
        from .db_utils import db_pool as pool
        
        if not pool:
            return {"error": "Database not connected"}
            
        async with pool.acquire() as conn:
            # Count total documents
            doc_count = await conn.fetchval("SELECT COUNT(*) FROM documents")
            
            # Count chunks
            chunk_count = await conn.fetchval("SELECT COUNT(*) FROM chunks")
            
            # Get sample document titles
            sample_docs = await conn.fetch("""
                SELECT title, source, LENGTH(content) as content_length 
                FROM documents 
                ORDER BY created_at DESC 
                LIMIT 5
            """)
            
            return {
                "document_count": doc_count,
                "chunk_count": chunk_count,
                "sample_documents": [
                    {
                        "title": doc["title"],
                        "source": doc["source"],
                        "content_length": doc["content_length"]
                    } for doc in sample_docs
                ],
                "question_generator_available": question_generator is not None
            }
    except Exception as e:
        logger.error(f"Debug endpoint error: {e}")
        return {"error": str(e)}


# Exception handlers
@app.exception_handler(Exception)
async def global_exception_handler(request: Request, exc: Exception):
    """Global exception handler."""
    logger.error(f"Unhandled exception: {exc}")
    
    return ErrorResponse(
        error=str(exc),
        error_type=type(exc).__name__,
        request_id=str(uuid.uuid4())
    )


# Development server
if __name__ == "__main__":
    uvicorn.run(
        app,
        host=APP_HOST,
        port=APP_PORT,
        reload=APP_ENV == "development",
        log_level=LOG_LEVEL.lower(),
    )